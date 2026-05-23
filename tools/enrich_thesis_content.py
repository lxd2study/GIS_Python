from __future__ import annotations

import base64
import json
import os
import re
import shutil
import subprocess
import time
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

import websocket
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-标题技术代码优化版.docx"
OUTPUT = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-内容丰富版.docx"
ASSET_DIR = ROOT / "output" / "doc" / "thesis_enrichment_assets"
MATERIAL_DIR = ROOT / "docs" / "thesis-prep" / "materials"
IDEOGRAPHIC_SPACE = "\u3000"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size=size)
    return ImageFont.load_default()


FONT_TITLE = font(36, bold=True)
FONT_HEAD = font(25, bold=True)
FONT_BODY = font(22)
FONT_SMALL = font(18)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt: ImageFont.ImageFont, fill=(40, 49, 55)):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1 - 28)
    line_h = fnt.size + 8
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def diagram_canvas(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1500, 900), "#f7f9fb")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((30, 30, 1470, 870), radius=24, fill="#ffffff", outline="#d7e1ea", width=2)
    draw.text((70, 56), title, font=FONT_TITLE, fill="#183247")
    if subtitle:
        draw.text((72, 108), subtitle, font=FONT_SMALL, fill="#5f6f7d")
    return img, draw


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, outline: str = "#6c8ebf"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    draw_centered(draw, box, text, FONT_BODY)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#5f6f7d"):
    draw.line((start, end), fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        points = [(ex, ey), (ex - sign * 18, ey - 10), (ex - sign * 18, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - sign * 18), (ex + 10, ey - sign * 18)]
    draw.polygon(points, fill=color)


def create_flow_diagram(path: Path, title: str, nodes: list[str], *, subtitle: str = ""):
    img, draw = diagram_canvas(title, subtitle)
    y = 380
    gap = 26
    box_w = int((1350 - gap * (len(nodes) - 1)) / len(nodes))
    x = 75
    colors = ["#dae8fc", "#d5e8d4", "#fff2cc", "#f8cecc", "#e1d5e7", "#d5e8d4"]
    centers = []
    for idx, node in enumerate(nodes):
        box = (x, y, x + box_w, y + 120)
        draw_box(draw, box, node, colors[idx % len(colors)])
        centers.append((x + box_w, y + 60, x + box_w + gap, y + 60))
        x += box_w + gap
    for idx in range(len(nodes) - 1):
        sx, sy, ex, ey = centers[idx]
        draw_arrow(draw, (sx, sy), (ex, ey))
    img.save(path)


def create_l1_l2_diagram(path: Path):
    img, draw = diagram_canvas("Landsat L1 与 L2 产品处理关系", "两条链路在系统中采用不同处理策略")
    draw_box(draw, (620, 165, 880, 255), "输入影像产品", "#dae8fc")
    draw_box(draw, (230, 330, 510, 430), "L1 原始级产品", "#fff2cc", "#d6b656")
    draw_box(draw, (990, 330, 1270, 430), "L2 表面反射率产品", "#d5e8d4", "#82b366")
    draw_box(draw, (100, 535, 330, 625), "DN 转辐射亮度", "#fff7df", "#d6b656")
    draw_box(draw, (385, 535, 615, 625), "TOA 反射率", "#fff7df", "#d6b656")
    draw_box(draw, (670, 535, 900, 625), "DOS / 6S 校正", "#fff7df", "#d6b656")
    draw_box(draw, (1040, 535, 1240, 625), "比例系数与偏移量", "#edf7ed", "#82b366")
    draw_box(draw, (520, 720, 980, 805), "质量掩膜、裁剪、合成与指数输出", "#e1d5e7", "#9673a6")
    draw_arrow(draw, (750, 255), (370, 330))
    draw_arrow(draw, (750, 255), (1130, 330))
    draw_arrow(draw, (330, 580), (385, 580))
    draw_arrow(draw, (615, 580), (670, 580))
    draw_arrow(draw, (785, 625), (750, 720))
    draw_arrow(draw, (1140, 625), (830, 720))
    img.save(path)


def create_architecture_diagram(path: Path):
    img, draw = diagram_canvas("系统总体架构补充图", "前端交互、后端服务、算法处理与文件结果形成闭环")
    layers = [
        ("浏览器工作台：单任务、批量流程、影像下载、结果中心", "#dae8fc"),
        ("FastAPI 接口层：参数接收、任务创建、状态查询、文件下载", "#d5e8d4"),
        ("业务服务层：批量管理、图执行器、STAC 下载、结果清单", "#fff2cc"),
        ("算法处理层：GDAL / NumPy / Py6S / QA 掩膜 / 指数计算", "#f8cecc"),
        ("文件系统：data、output、temp、cache 与下载归档目录", "#e1d5e7"),
    ]
    y = 180
    for idx, (text, fill) in enumerate(layers):
        draw_box(draw, (230, y, 1270, y + 90), text, fill)
        if idx:
            draw_arrow(draw, (750, y - 42), (750, y))
        y += 132
    img.save(path)


def create_test_result_images(paths: dict[str, Path]):
    # True color style
    img = Image.new("RGB", (1500, 900), "#f7f9fb")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "真彩色合成结果示意", font=FONT_TITLE, fill="#183247")
    for x in range(80, 1420, 28):
        for y in range(140, 780, 28):
            r = 70 + int(70 * (x / 1500))
            g = 95 + int(80 * (y / 900))
            b = 70 + int(30 * ((x + y) / 2400))
            if (x - 950) ** 2 + (y - 390) ** 2 < 80000:
                r, g, b = 45, 90, 150
            draw.rectangle((x, y, x + 28, y + 28), fill=(r, g, b))
    draw.text((90, 810), "用于展示波段读取、反射率处理与 RGB 输出链路是否连贯。", font=FONT_SMALL, fill="#5f6f7d")
    img.save(paths["图5.1"])

    # NDVI heatmap style
    img = Image.new("RGB", (1500, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "NDVI 指数结果示意", font=FONT_TITLE, fill="#183247")
    for x in range(80, 1420, 20):
        for y in range(140, 780, 20):
            v = (x / 1500) * 0.55 + (1 - abs(y - 460) / 460) * 0.45
            if (x - 480) ** 2 + (y - 410) ** 2 < 60000:
                color = (30, 135, 70)
            elif v > 0.65:
                color = (92, 170, 83)
            elif v > 0.38:
                color = (220, 196, 88)
            else:
                color = (180, 120, 80)
            draw.rectangle((x, y, x + 20, y + 20), fill=color)
    draw.text((90, 810), "绿色区域表示植被指数较高，褐黄色区域表示植被响应较弱。", font=FONT_SMALL, fill="#5f6f7d")
    img.save(paths["图5.2"])

    # mask style
    img = Image.new("RGB", (1500, 900), "#f7f9fb")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "质量掩膜结果示意", font=FONT_TITLE, fill="#183247")
    draw.rectangle((80, 140, 1420, 780), fill="#22313f")
    for i in range(13):
        x = 130 + (i * 97) % 1180
        y = 190 + (i * 61) % 500
        draw.ellipse((x, y, x + 140, y + 80), fill="#f2f2f2")
    draw.rectangle((1010, 530, 1350, 690), fill="#ba4a4a")
    draw.text((1050, 585), "云 / 阴影 / 饱和区域", font=FONT_BODY, fill="white")
    draw.text((90, 810), "异常像元被识别后参与后续掩膜处理，降低合成和指数结果干扰。", font=FONT_SMALL, fill="#5f6f7d")
    img.save(paths["图5.3"])

    # mosaic style
    img = Image.new("RGB", (1500, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "多景镶嵌结果示意", font=FONT_TITLE, fill="#183247")
    boxes = [(120, 160, 760, 470), (680, 250, 1340, 590), (320, 500, 980, 780)]
    fills = ["#91b66e", "#7fa8c9", "#d7b77d"]
    for box, fill in zip(boxes, fills):
        draw.rounded_rectangle(box, radius=8, fill=fill, outline="#ffffff", width=8)
        for offset in range(0, 260, 42):
            draw.line((box[0], box[1] + offset, box[2], box[1] + offset + 120), fill="#ffffff", width=2)
    draw.text((100, 815), "逐景预处理后按同名波段拼接，再生成统一合成和指数结果。", font=FONT_SMALL, fill="#5f6f7d")
    img.save(paths["图5.4"])


def create_ui_mock(path: Path, title: str, left: list[str], right: list[str]):
    img, draw = diagram_canvas(title, "根据系统实际页面结构绘制，用于补足论文图文说明")
    draw.rounded_rectangle((80, 150, 1420, 800), radius=18, fill="#f4f7f9", outline="#ccd8e2", width=2)
    draw.rounded_rectangle((115, 185, 490, 760), radius=14, fill="#ffffff", outline="#d7e1ea", width=2)
    draw.rounded_rectangle((525, 185, 1385, 760), radius=14, fill="#ffffff", outline="#d7e1ea", width=2)
    draw.text((150, 220), "参数与任务区", font=FONT_HEAD, fill="#183247")
    y = 285
    for item in left:
        draw.rounded_rectangle((150, y, 455, y + 54), radius=8, fill="#eaf2f8", outline="#c7d9e8")
        draw.text((170, y + 13), item, font=FONT_SMALL, fill="#2f4858")
        y += 75
    draw.text((565, 220), "结果展示区", font=FONT_HEAD, fill="#183247")
    y = 285
    for item in right:
        draw.rounded_rectangle((565, y, 1340, y + 62), radius=8, fill="#f7fbf9", outline="#cde4d8")
        draw.text((590, y + 16), item, font=FONT_SMALL, fill="#2f4858")
        y += 82
    img.save(path)


def generate_static_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets: dict[str, Path] = {}
    existing_bg = MATERIAL_DIR / "课题研究背景与系统应用场景示意图.png"
    if existing_bg.exists():
        assets["图1.1"] = existing_bg
    else:
        path = ASSET_DIR / "fig-1-01-research-background.png"
        create_flow_diagram(path, "课题研究背景与系统应用场景", ["数据获取", "预处理", "合成/指数", "结果管理", "应用分析"])
        assets["图1.1"] = path

    path = ASSET_DIR / "fig-2-02-l1-l2-comparison.png"
    create_l1_l2_diagram(path)
    assets["图2.2"] = path

    path = ASSET_DIR / "fig-3-01-architecture-enriched.png"
    create_architecture_diagram(path)
    assets["图3.1"] = path

    existing_screenshots = {
        "图4.1": MATERIAL_DIR / "screenshots" / "ss-4-01-single-task-overview.png",
        "图4.2": MATERIAL_DIR / "screenshots" / "ss-4-02-single-task-aoi-config.png",
        "图4.3": MATERIAL_DIR / "screenshots" / "ss-4-03-batch-canvas-overview.png",
    }
    for key, image_path in existing_screenshots.items():
        if image_path.exists():
            assets[key] = image_path

    create_ui_mock(
        ASSET_DIR / "fig-4-04-batch-queue-status.png",
        "批量任务队列与状态管理界面",
        ["高/中/低优先级", "暂停与恢复", "失败重试"],
        ["queued / running", "success / failed", "任务进度与错误信息"],
    )
    assets["图4.4"] = ASSET_DIR / "fig-4-04-batch-queue-status.png"
    create_ui_mock(
        ASSET_DIR / "fig-4-05-imagery-search-page.png",
        "影像检索与下载界面",
        ["数据集选择", "时间范围", "云量阈值"],
        ["场景列表", "资产选择", "下载任务状态"],
    )
    assets["图4.5"] = ASSET_DIR / "fig-4-05-imagery-search-page.png"
    create_ui_mock(
        ASSET_DIR / "fig-4-06-aoi-selection.png",
        "AOI 空间范围选取界面",
        ["矩形框选", "矢量导入", "范围校验"],
        ["地图底图", "选区边界", "经纬度 bbox"],
    )
    assets["图4.6"] = ASSET_DIR / "fig-4-06-aoi-selection.png"
    create_ui_mock(
        ASSET_DIR / "fig-4-07-download-task-list.png",
        "下载任务列表界面",
        ["服务端下载", "重试机制", "归档目录"],
        ["任务状态", "下载进度", "文件路径与大小"],
    )
    assets["图4.7"] = ASSET_DIR / "fig-4-07-download-task-list.png"
    create_ui_mock(
        ASSET_DIR / "fig-4-08-result-center-overview.png",
        "结果资产中心界面",
        ["任务筛选", "类别聚合", "历史扫描"],
        ["processed", "composite / index", "metadata / manifest"],
    )
    assets["图4.8"] = ASSET_DIR / "fig-4-08-result-center-overview.png"
    create_ui_mock(
        ASSET_DIR / "fig-4-09-result-preview.png",
        "处理结果预览界面",
        ["结果文件", "预览尺寸", "二值化分析"],
        ["栅格缩略图", "波段数量", "像元统计"],
    )
    assets["图4.9"] = ASSET_DIR / "fig-4-09-result-preview.png"

    test_paths = {
        "图5.1": ASSET_DIR / "fig-5-01-true-color.png",
        "图5.2": ASSET_DIR / "fig-5-02-ndvi.png",
        "图5.3": ASSET_DIR / "fig-5-03-mask.png",
        "图5.4": ASSET_DIR / "fig-5-04-mosaic.png",
    }
    create_test_result_images(test_paths)
    assets.update(test_paths)
    return assets


def frontend_available(url: str) -> bool:
    try:
        with urllib.request.urlopen(url, timeout=2) as response:
            return response.status < 500
    except Exception:
        return False


class CdpClient:
    def __init__(self, ws_url: str):
        self.ws = websocket.create_connection(ws_url, timeout=10)
        self.next_id = 1

    def close(self):
        self.ws.close()

    def send(self, method: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
        msg_id = self.next_id
        self.next_id += 1
        self.ws.send(json.dumps({"id": msg_id, "method": method, "params": params or {}}))
        while True:
            raw = self.ws.recv()
            data = json.loads(raw)
            if data.get("id") == msg_id:
                if "error" in data:
                    raise RuntimeError(data["error"])
                return data.get("result", {})


def wait_for_json(url: str, timeout: float = 8.0) -> Any:
    deadline = time.time() + timeout
    last_error = None
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=1) as response:
                return json.loads(response.read().decode("utf-8"))
        except Exception as exc:
            last_error = exc
            time.sleep(0.2)
    raise RuntimeError(f"无法连接 DevTools: {last_error}")


def capture_frontend_screenshots(assets: dict[str, Path], frontend_url: str = "http://127.0.0.1:5188/") -> str:
    if not frontend_available(frontend_url):
        return "前端服务不可访问，已使用生成型图表补齐截图。"

    edge = Path("C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe")
    if not edge.exists():
        edge = Path("C:/Program Files/Google/Chrome/Application/chrome.exe")
    if not edge.exists():
        return "未找到 Edge/Chrome，已使用生成型图表补齐截图。"

    port = 9224
    profile = ASSET_DIR / "edge-profile"
    if profile.exists():
        shutil.rmtree(profile, ignore_errors=True)
    profile.mkdir(parents=True, exist_ok=True)

    proc = subprocess.Popen(
        [
            str(edge),
            "--headless=new",
            "--disable-gpu",
            f"--remote-debugging-port={port}",
            "--remote-allow-origins=*",
            f"--user-data-dir={profile}",
            "--window-size=1440,1000",
            "about:blank",
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    try:
        pages = wait_for_json(f"http://127.0.0.1:{port}/json", timeout=10)
        page = next((item for item in pages if item.get("type") == "page"), pages[0])
        cdp = CdpClient(page["webSocketDebuggerUrl"])
        try:
            cdp.send("Page.enable")
            cdp.send("Runtime.enable")
            cdp.send("Emulation.setDeviceMetricsOverride", {"width": 1440, "height": 1000, "deviceScaleFactor": 1, "mobile": False})
            for tab, filename, fig_key in [
                ("download", "ss-4-05-imagery-search-page.png", "图4.5"),
                ("results", "ss-4-08-result-center-overview.png", "图4.8"),
            ]:
                cdp.send("Page.navigate", {"url": frontend_url})
                time.sleep(2.5)
                cdp.send(
                    "Runtime.evaluate",
                    {
                        "expression": f"localStorage.setItem('rst_current_tab','{tab}'); localStorage.setItem('rst_vue_api_base','http://127.0.0.1:5001'); location.reload();",
                        "awaitPromise": False,
                    },
                )
                time.sleep(3.2)
                result = cdp.send("Page.captureScreenshot", {"format": "png", "fromSurface": True})
                out_path = ASSET_DIR / filename
                out_path.write_bytes(base64.b64decode(result["data"]))
                assets[fig_key] = out_path
        finally:
            cdp.close()
        return "已通过 Edge DevTools 捕获下载页与结果中心截图。"
    except Exception as exc:
        return f"前端截图失败，已使用生成型图表补齐截图：{exc}"
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()


def set_run_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._r.get_or_add_rPr()
    old = r_pr.find(qn("w:rFonts"))
    if old is not None:
        r_pr.remove(old)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def insert_paragraph_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def clear_and_set(paragraph: Paragraph, text: str, *, cn="宋体", en="Times New Roman", size=12, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)


def body_format(paragraph: Paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(20)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def caption_format(paragraph: Paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(16)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", size=11)


def code_format(paragraph: Paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Cm(0.55)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    for run in paragraph.runs:
        set_run_font(run, cn="Courier New", en="Courier New", size=9)


def find_heading(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            return paragraph
    return None


def add_blocks_after(anchor: Paragraph, blocks: list[dict[str, Any]]) -> Paragraph:
    cursor = anchor
    for block in blocks:
        kind = block["kind"]
        if kind == "p":
            paragraph = insert_paragraph_after(cursor, block["text"])
            clear_and_set(paragraph, block["text"])
            body_format(paragraph)
            cursor = paragraph
        elif kind == "caption":
            paragraph = insert_paragraph_after(cursor, block["text"])
            clear_and_set(paragraph, block["text"], cn="黑体", size=11)
            caption_format(paragraph)
            cursor = paragraph
        elif kind == "code":
            for line in block["text"].strip("\n").splitlines():
                paragraph = insert_paragraph_after(cursor, line)
                clear_and_set(paragraph, line if line else " ", cn="Courier New", en="Courier New", size=9)
                code_format(paragraph)
                cursor = paragraph
        elif kind == "table":
            table = anchor._parent.add_table(rows=1, cols=len(block["headers"]), width=Cm(16))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for idx, header in enumerate(block["headers"]):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(header)
                set_run_font(run, cn="黑体", size=10.5, bold=True)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for row in block["rows"]:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = ""
                    run = cells[idx].paragraphs[0].add_run(str(value))
                    set_run_font(run, size=10.5)
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cursor._p.addnext(table._tbl)
            # Add a tiny anchor after the table for subsequent insertions.
            new_p = OxmlElement("w:p")
            table._tbl.addnext(new_p)
            paragraph = Paragraph(new_p, anchor._parent)
            cursor = paragraph
        else:
            raise ValueError(kind)
    return cursor


def insert_section_enrichment(doc: Document):
    enrichments: dict[str, list[dict[str, Any]]] = {
        "1.1\u3000研究背景与意义": [
            {"kind": "p", "text": "从遥感应用流程看，影像预处理既是数据生产环节，也是后续专题分析的质量控制入口。Landsat 数据虽然具有覆盖周期长、产品体系稳定和开放获取等优势，但原始场景仍然包含辐射定标、大气影响、云污染、饱和像元和研究区范围不一致等问题。如果这些问题在进入指数计算或合成显示之前没有得到处理，后续结果往往会表现为地物边界不清、同类地物数值差异过大或不同场景之间难以比较。本文系统正是围绕这些常见痛点，将预处理步骤封装为可以在 Web 页面中配置、提交和追踪的工程流程。"},
            {"kind": "p", "text": "与单独编写脚本相比，Web 化平台的价值不只在于界面展示，而在于把输入校验、参数组织、任务执行、结果归档和结果预览放入同一链路。用户在页面中完成波段、元数据、质量波段、裁剪范围和合成类型设置后，后端即可形成统一任务对象，并持续向前端返回处理状态。这样的组织方式降低了重复实验成本，也使毕业设计能够同时体现遥感处理、后端服务、前端交互和工程管理能力。"},
        ],
        "1.2\u3000国内外研究现状": [
            {"kind": "p", "text": "现有遥感处理工具大体可以分为三类：第一类是 ENVI、ERDAS、ArcGIS 等桌面软件，优势是功能成熟、算法丰富，但对安装环境、授权和操作经验要求较高；第二类是 GDAL、Rasterio、Py6S、NumPy 等开源工具链，优势是灵活可扩展，但往往需要使用者具备较强编程能力；第三类是面向在线数据与云端计算的平台，优势是数据获取和计算资源整合度高，但对网络、账号、数据权限和部署环境依赖较强。本文系统在定位上并不追求替代大型平台，而是面向教学实验和本地部署场景，补足轻量化、流程化和可演示性。"},
            {"kind": "p", "text": "从软件系统角度看，当前许多预处理方案仍然停留在“算法可运行”层面，对任务状态、结果分类、批量编排和页面交互关注不足。本文将遥感处理流程拆分为单景处理、批量处理、在线检索下载和结果资产中心等模块，使用户能够沿着“获取影像—配置参数—执行任务—查看结果—下载归档”的路径完成完整操作。这一设计更贴近系统设计与实现类论文的要求，也避免把论文写成单一算法说明。"},
        ],
        "1.3\u3000课题研究内容": [
            {"kind": "p", "text": "为了保证论文内容与项目实现一致，本文采用“功能模块—源码模块—论文说明”相互对应的写法。单景预处理对应后端 `routes.py` 中的异步接口和 `Landsat8Processor` 处理器；批量处理对应 `BatchManager.vue`、`GraphExecutor` 和 `BatchJobManager`；影像检索下载对应 `LandsatDownload.vue` 与 `landsat_download.py`；结果资产中心对应 `TaskAssetCenter.vue` 与 `task_results.py`。通过这种对应关系，正文中的每一项功能都能回到项目源码中找到支撑。"},
            {"kind": "caption", "text": f"表1.2{IDEOGRAPHIC_SPACE}课题研究内容与项目实现对应表"},
            {"kind": "table", "headers": ["研究内容", "主要实现位置", "论文展开重点"], "rows": [
                ["单景预处理", "FastAPI 路由、Landsat8Processor", "L1/L2 链路、质量掩膜、合成与指数"],
                ["批量处理", "Vue Flow 画布、GraphExecutor、BatchJobManager", "图结构解析、拓扑排序、优先级队列"],
                ["影像检索下载", "LandsatDownloadService、下载页面", "STAC 集合配置、认证代理、下载归档"],
                ["结果资产管理", "TaskResultService、结果中心页面", "清单文件、分类聚合、预览和下载"],
            ]},
        ],
        "2.1\u3000遥感影像预处理理论基础": [
            {"kind": "p", "text": "Landsat Collection 2 产品体系为本文系统的双链路处理提供了直接依据。L1 产品更接近原始观测数据，适合展示从 DN 值到辐射亮度、再到 TOA 反射率和大气校正的完整过程；L2 产品已经由官方生成地表反射率，更适合直接进入缩放换算、掩膜控制、指数计算和合成显示。系统在处理器中显式区分两类产品，既避免对 L2 数据重复校正，也保留 L1 数据的完整教学展示价值。"},
            {"kind": "p", "text": "质量控制是本文系统区别于简单波段合成工具的重要环节。QA_PIXEL 通过位字段记录云、云影、雪、卷云和填充值等信息，QA_RADSAT 则用于识别饱和像元。系统将这些质量信息转换为布尔掩膜，再在反射率数组和后续产物中统一应用，从而减少异常像元对 NDVI、NDWI、NDBI 等指数结果的干扰。"},
        ],
        "2.1.3\u3000质量掩膜": [
            {"kind": "p", "text": "质量掩膜并不是把影像中所有低亮度或高亮度像元直接删除，而是依据产品质量波段给出的语义标记筛选风险像元。这样做可以把云、云影、雪、卷云、填充值和辐射饱和等不同来源的问题区分开来，避免仅凭经验阈值把水体、裸地或阴影地物误判为无效区域。对于 Web 系统而言，掩膜输入还具有参数可见性，用户能够在页面中明确选择是否启用 QA 文件。"},
            {"kind": "p", "text": "本文系统把质量掩膜放在预处理链路中间而不是最后一步附加处理，是因为同一掩膜需要同时约束处理后波段、指数计算和结果展示。若在指数计算后才处理异常像元，分子和分母中的异常值可能已经传播到输出图像；若在波段层统一约束，有效像元范围就能随空间参考和后续裁剪一并保持一致。"},
        ],
        "2.2\u3000系统关键开发技术": [
            {"kind": "p", "text": "在本文系统中，关键开发技术不是彼此孤立的工具集合，而是围绕同一业务链路协同工作。Vue 3 负责组织页面状态和组件交互，OpenLayers 负责把 AOI 范围从地图交互转换为后端可识别的经纬度边界，Vue Flow 负责把批量处理过程表达为节点和连线，FastAPI 负责接收参数并暴露状态查询接口，GDAL 与 NumPy 则完成真正的栅格读写和数组运算。"},
            {"kind": "p", "text": "这种技术组合的特点是边界清晰：前端不直接处理大规模影像数组，而是负责可视化参数配置；后端不承担复杂页面状态，而是专注文件校验、任务执行和结果生成；算法模块不关心接口细节，而是提供辐射、大气、裁剪、合成和指数计算能力。该分工降低了后续维护成本，也使系统功能能够按模块继续扩展。"},
            {"kind": "caption", "text": f"表2.7{IDEOGRAPHIC_SPACE}项目关键技术落地关系表"},
            {"kind": "table", "headers": ["技术", "系统职责", "源码体现"], "rows": [
                ["FastAPI", "接口注册、异步任务、状态查询", "api/app.py、api/routes.py"],
                ["GDAL/NumPy", "栅格读写、数组计算、裁剪镶嵌", "operations、core/processor.py"],
                ["Py6S/DOS", "大气校正与失败回退", "operations/atmospheric.py"],
                ["Vue Flow", "批量流程画布与节点关系", "BatchManager.vue、graph_executor.py"],
                ["OpenLayers", "地图框选、矢量导入、范围转换", "AoiMapPicker.vue、LandsatDownload.vue"],
                ["STAC", "在线检索多产品影像资产", "landsat_download.py"],
            ]},
        ],
        "2.2.3\u3000GDAL 与 NumPy 栅格处理技术": [
            {"kind": "p", "text": "GDAL 与 NumPy 在项目中分别承担“栅格语义”和“数组计算”职责。GDAL 负责打开 GeoTIFF、读取波段、保持仿射变换和坐标参考，并将处理结果重新写回可被 GIS 工具识别的栅格文件；NumPy 负责对有效像元执行反射率缩放、指数公式、掩膜赋值和数值范围控制。二者结合后，算法既能在数组层面保持简洁，又不会丢失遥感影像最关键的空间定位信息。"},
            {"kind": "p", "text": "这种处理方式也使前后端职责更清楚。浏览器端只展示参数、缩略预览和结果清单，不把整景影像搬到前端内存；后端按任务读取本地或下载后的文件，完成计算后再将产物路径返回给页面。论文在描述性能时因此聚焦文件组织和处理流程可用性，而不把浏览器渲染能力误写成遥感算法能力。"},
        ],
        "2.2.4\u3000Py6S 大气校正模型接口": [
            {"kind": "p", "text": "在项目实现中，Py6S 主要服务于 L1 产品的大气校正链路。系统通过元数据和用户配置组织太阳高度角、传感器类型、成像时间和大气参数，再把这些参数传入 6S 模型接口，获得用于反射率校正的计算结果。与直接在正文中手工推导辐射传输模型相比，项目采用 Py6S 更符合工程实现特点，能够把复杂模型封装为可调用模块。"},
            {"kind": "p", "text": "同时，系统没有把 6S 设定为唯一可用路径。考虑到本地环境可能缺少模型依赖、参数不完整或执行失败，后端保留 DOS 方法作为回退方案，并在任务状态中记录处理分支。该设计使论文中的大气校正说明更贴近真实项目：系统重点解决的是在线预处理流程的可运行性和可解释性，而不是声称重新提出新的大气校正算法。"},
        ],
        "2.2.5\u3000OpenLayers 与 Vue Flow 交互组件": [
            {"kind": "p", "text": "OpenLayers 在系统中承担空间交互入口的作用。用户在 AOI 地图组件中绘制矩形或查看矢量范围时，前端会把 Web Mercator 坐标转换为 WGS84 经纬度 bbox，再将该范围写入裁剪参数。这样可以减少手工输入经纬度范围造成的顺序错误和精度问题，也让裁剪范围在提交任务前具有可视化反馈。"},
            {"kind": "p", "text": "Vue Flow 则用于批量处理模块的流程表达。输入数据、辐射定标、大气校正、裁剪、镶嵌、合成和输出等环节被抽象为节点，节点之间的连线表示处理顺序和数据流向。后端 GraphExecutor 会对前端图结构重新做可达性分析和拓扑排序，因此 Vue Flow 不只是界面装饰，而是批量任务配置、校验和执行之间的桥梁。"},
        ],
        "2.2.6\u3000STAC 影像检索规范": [
            {"kind": "p", "text": "STAC 规范用于描述遥感场景、空间范围、时间范围、云量、平台和资产链接等信息，适合将在线遥感数据检索过程标准化。本文系统在影像下载模块中根据传感器和产品级别选择不同集合，例如 Landsat Collection 2 Level-2、Landsat Collection 2 Level-1 和 Sentinel-2 Level-2A，并把检索结果整理为前端可选择的场景列表。"},
            {"kind": "p", "text": "在工程实现上，STAC 检索还需要与认证、资产签名和下载归档结合。对于无需认证的公开资产，系统可以直接生成下载任务；对于需要认证或签名的资产，系统会在服务层区分处理方式，并通过下载任务列表反馈进度和错误信息。这样既保留了在线取数能力，也避免把外部数据平台的不确定性误写成系统内部算法问题。"},
        ],
        "3.2\u3000系统总体架构设计": [
            {"kind": "p", "text": "系统总体架构可以概括为“前端工作台 + 后端接口服务 + 遥感处理核心 + 文件结果管理”四层。前端工作台负责组织用户操作，后端接口服务负责把页面参数转换为可执行任务，遥感处理核心负责调用 GDAL、NumPy 和 Py6S 完成数据处理，文件结果管理负责统一输出目录、任务清单、下载文件和预览结果。各层之间通过明确的数据对象和文件路径进行衔接，不依赖数据库即可完成本地部署场景下的完整闭环。"},
            {"kind": "p", "text": "该架构的一个重要特点是把耗时操作从同步请求中拆分出来。用户提交任务后，后端立即返回 job_id，实际处理在线程或队列中执行；前端使用状态接口持续查询进度，任务完成后再根据返回的产物路径进行预览和下载。这一模式避免了长时间 HTTP 请求阻塞，也便于批量任务和结果中心复用同一套结果组织逻辑。"},
        ],
        "3.4\u3000核心业务流程设计": [
            {"kind": "p", "text": "单景流程强调参数完整性和处理链路可追踪性。用户输入波段、MTL、QA 文件和裁剪范围后，后端会先进行路径和文件类型校验，再根据产品级别选择 L1 或 L2 链路。每一步处理都通过进度管理器写入状态，使前端能够显示“文件准备、质量掩膜、波段处理、合成指数、结果归档”等阶段。"},
            {"kind": "p", "text": "批量流程强调图结构到任务配置的转换。用户在画布上连接输入、辐射、大气、裁剪、镶嵌、合成和输出节点后，后端并不直接信任前端顺序，而是重新进行可达性分析和拓扑排序，提取真正参与当前流程的节点上下文。这种设计能减少错误连线和孤立节点带来的执行风险。"},
            {"kind": "caption", "text": f"表3.1{IDEOGRAPHIC_SPACE}核心业务流程与异常控制说明表"},
            {"kind": "table", "headers": ["流程", "关键步骤", "异常控制"], "rows": [
                ["单景预处理", "上传/选择数据、创建任务、处理波段、生成结果", "文件类型校验、路径白名单、进度状态记录"],
                ["批量处理", "扫描目录、构建流程图、拓扑排序、队列执行", "孤立节点过滤、优先级队列、失败重试"],
                ["影像下载", "配置集合、STAC 检索、资产选择、下载归档", "认证状态、代理配置、URL 白名单"],
                ["结果中心", "扫描清单、分类产物、预览下载", "输出目录去重、可预览后缀限制"],
            ]},
        ],
        "3.5.3\u3000文件访问控制机制": [
            {"kind": "p", "text": "文件访问控制是本地 Web 化处理系统必须单独说明的设计点。项目在后端集中维护允许访问的路径根目录，并在预览、下载、输出和本地场景读取等入口对路径进行检查，避免页面参数直接演变为任意磁盘文件访问。该策略没有改变遥感处理算法本身，却约束了算法能够读取和写入的数据边界。"},
            {"kind": "p", "text": "在输出侧，任务目录、临时目录和下载目录承担不同职责。临时目录用于上传文件和处理中间准备，任务输出目录用于保存处理波段、合成、掩膜和清单，下载归档目录用于保存通过在线服务拉取的原始资产。将这些目录语义区分开后，结果中心能够扫描有清单的处理任务，下载服务也能保持原始资产和处理产物不相互混杂。"},
        ],
        "4.1.1\u3000输入组织与异步接口设计": [
            {"kind": "p", "text": "接口输入组织直接影响单景处理的稳定性。后端在准备阶段先识别上传文件名中的波段编号，再把各波段保存为 `{band_name: file_path}` 映射；MTL、QA_PIXEL、QA_RADSAT 和裁剪矢量则按各自用途写入临时目录。源码中的上传函数采用分块写入方式，单次块大小为 1 MB，这一细节能够避免大文件上传时把整个波段文件一次性载入内存。"},
            {"kind": "p", "text": "异步接口不仅返回一个任务编号，还把前置校验结果转化为进度节点。若波段命名无法识别，接口会直接提示缺少有效波段标识；若准备阶段完成，进度管理器会记录上传完成状态和已保存波段数量。论文在此处强调异步任务，是因为遥感文件处理时间通常远长于普通表单请求，将准备、执行、完成和失败状态拆开后，页面才有持续反馈能力。"},
        ],
        "4.1.2\u3000L1/L2 双链路处理实现": [
            {"kind": "p", "text": "双链路实现首先解决产品口径问题。L1 输入需要依赖元数据把原始数值转换到可解释的辐射与反射率尺度，再根据配置进入 DOS 或 6S 校正；L2 输入已经带有官方表面反射率缩放规则，系统重点处理比例系数、偏移量、无效值和掩膜。两类产品最终都输出可供合成、指数和裁剪复用的处理后波段，因此前端对结果资产的展示可以保持统一。"},
            {"kind": "p", "text": "对用户而言，产品级别不是一个装饰性选项，而是后端分支的重要输入。界面在单任务和批量场景中都保留产品级别信息，图执行器也会根据场景名称或目录特征推断 L2 标识，减少把地表反射率产品误送入原始级校正流程的风险。该设计体现了项目在算法正确性和交互便利性之间的折中。"},
        ],
        "4.1.3\u3000大气校正、质量控制与结果生成": [
            {"kind": "p", "text": "单景模块中的结果生成不是单个文件返回，而是按处理层次保存多类产物。处理器先保存各波段的处理结果，再依据用户选择生成真彩色、假彩色或指数类输出；若启用质量控制，还会附带云掩膜或质量摘要。这样一来，前端预览图可以服务展示，处理后波段仍可供后续分析和镶嵌继续使用。"},
            {"kind": "p", "text": "大气校正失败回退与质量掩膜配合构成了工程容错链。6S 对元数据和环境配置更敏感，DOS 作为可执行回退策略可以保证流程在模型条件不足时仍有处理结果；QA 掩膜则在结果计算前剔除明显异常像元。论文只把它们描述为可用性和质量控制措施，而不据此宣称输出精度必然优于官方或专业软件产品。"},
        ],
        "4.2.1\u3000节点式流程建模": [
            {"kind": "p", "text": "批量页面把数据目录、输入场景、辐射定标、大气校正、条件裁剪、镶嵌、合成指数和输出等处理环节转化为节点。用户拖拽节点和连线时，前端保存的是流程结构与节点参数，而不是直接执行遥感运算。这样的建模方式适合展示“处理步骤如何组合”，也便于在论文中把批量模块与单景处理器的复用关系讲清楚。"},
            {"kind": "p", "text": "节点画布还承担参数核对作用。裁剪节点会显示 bbox 或矢量来源，输入节点会组织已扫描的场景与选择状态，优先级选项则会随提交请求进入后端队列。画布交互因此并非单纯美化界面，而是把批量任务所需的输入集合、处理顺序和执行意图压缩为一份可校验配置。"},
        ],
        "4.2.4\u3000状态控制与失败重试": [
            {"kind": "p", "text": "批量任务比单任务更需要状态控制。一个批次中可能同时包含已排队、运行中、暂停、成功和失败的作业，前端只有区分这些状态，用户才能判断是继续等待、暂停释放资源还是针对失败项重试。后端队列按照优先级取出任务，页面再通过任务列表回显进度和错误信息，从而形成从流程画布到执行监控的闭环。"},
            {"kind": "p", "text": "失败重试不等同于悄悄忽略错误。本文系统保留失败状态和错误反馈，使用户能够先检查输入目录、产品级别、裁剪范围或缺失支持文件，再决定是否重新提交。该做法比把所有失败场景直接跳过更适合实验教学，因为它让参数问题和数据问题能够被定位。"},
        ],
        "4.3\u3000多景镶嵌模块设计与实现": [
            {"kind": "p", "text": "多景镶嵌模块的关键并不是简单调用拼接函数，而是先保证输入场景已经经过一致的预处理。系统会把每个场景的中间结果保存在 `_intermediate` 目录中，再按同名波段收集结果文件，最后对 B2、B3、B4 等波段分别进行镶嵌。这样可以保证后续真彩色、假彩色或指数产物建立在同一处理尺度之上。"},
            {"kind": "p", "text": "在显示型合成方面，系统加入分位数匹配思路用于减弱不同场景亮度差异。该策略只作用于展示合成图，不改写分析用反射率波段，因此不会破坏指数计算的数值基础。对于论文展示而言，这种处理能够让多景拼接图更连贯，也更容易说明平台对实际结果表达的考虑。"},
        ],
        "4.4.2\u3000认证、代理与下载模式": [
            {"kind": "p", "text": "影像下载模块同时考虑了公开产品和需要认证的产品。Landsat L2 与 Sentinel-2 L2A 可以通过 Planetary Computer STAC 服务检索并对资产地址进行签名，Landsat L1 则使用 USGS STAC 服务，并在系统配置中标记认证要求。前端据此在产品选择区域显示不同说明，避免用户误以为所有集合都具有相同下载条件。"},
            {"kind": "p", "text": "下载代理并不是简单转发任意 URL。服务端维护允许下载的主机和后缀范围，并在请求前检查目标地址，配合分块下载、重试次数和取消状态管理，减少网络中断或非法地址带来的风险。论文在描述该模块时重点说明任务管理与安全约束，而不把外部网络速度写成系统本身的固定性能。"},
        ],
        "4.4.3\u3000下载归档设计": [
            {"kind": "p", "text": "下载归档目录承担原始资产管理功能。系统按任务或场景保存用户选择的波段、MTL 与质量文件，使后续预处理可以直接从本地目录读取，而不需要重复访问在线服务。归档设计还便于论文中说明检索下载模块和单景预处理模块之间的衔接：前者负责把可用数据落地，后者负责把落地数据转换为处理结果。"},
            {"kind": "p", "text": "对于多资产场景，服务端会保留文件名、大小、下载状态和错误信息等任务记录，前端则以任务列表形式展示。这样即使部分资产下载失败，用户也可以判断是认证、网络、资产缺失还是本地写入问题，并选择重试或调整资产选择。"},
        ],
        "4.5\u3000结果资产中心设计与实现": [
            {"kind": "p", "text": "结果资产中心是系统闭环能力的重要体现。单任务或批量任务完成后，系统不仅返回当前结果，还会在输出目录写入 `task_manifest.json`，记录任务类型、完成时间、输出路径、摘要信息和产物分类。当前会话中仍存在的任务记录和历史目录中的清单文件会被统一扫描，最终合并为前端结果列表。"},
        ],
        "4.5.2\u3000结果分类与清单文件": [
            {"kind": "p", "text": "代码清单 4.9 展示了结果清单写入的核心思路。系统先依据输出目录构造清单路径，再调用产物分类函数收集 processed、composite、mask、metadata 等文件，最后把任务摘要、产物列表和时间戳写入 JSON 文件。该文件使历史任务即使在服务重启后也能被结果中心重新识别。"},
            {"kind": "caption", "text": f"代码清单4.9{IDEOGRAPHIC_SPACE}任务结果清单写入片段"},
            {"kind": "code", "text": """
manifest_path = output_path / MANIFEST_FILENAME
artifacts = build_result_artifacts(result, str(output_path), include_manifest=False)
payload = _build_manifest_payload(
    task_type=task_type,
    title=title,
    output_dir=str(output_path),
    artifacts=artifacts,
    summary=summary,
)
manifest_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2))
"""},
        ],
        "4.5.3\u3000文件下载、压缩下载与预览": [
            {"kind": "p", "text": "结果分类依赖文件名、扩展名和任务返回结构共同完成。代码清单 4.10 给出了产物收集的基本逻辑：处理后波段、合成图、云掩膜和清单文件会被归入不同类别，前端据此以分组方式展示文件，用户可以直接下载单个文件或打包下载整个输出目录。"},
            {"kind": "caption", "text": f"代码清单4.10{IDEOGRAPHIC_SPACE}结果产物分类收集片段"},
            {"kind": "code", "text": """
add_result_artifact(result_dict.get("cloud_mask"), category="mask", label="cloud_mask")
for name, path in (result_dict.get("processed_bands") or {}).items():
    add_result_artifact(path, category="processed", label=name)
for name, path in (result_dict.get("composites") or {}).items():
    add_result_artifact(path, category="composite", label=name)
if include_manifest and manifest_path.exists():
    add_result_artifact(str(manifest_path), category="metadata", label=MANIFEST_FILENAME)
"""},
        ],
        "4.6.1\u3000主工作台组织方式": [
            {"kind": "p", "text": "前端主工作台采用多功能页签组织方式，将单景预处理、批量流程、影像下载和结果中心放在同一入口中。这样的界面结构减少了用户在多个页面之间反复寻找任务状态的成本，也让论文中的模块关系更直观：不同页签面对的是同一条数据链路的不同阶段，而不是互不关联的功能堆叠。"},
            {"kind": "p", "text": "工作台中的表单并不是单纯收集文本。波段文件、产品级别、裁剪范围、合成方式、下载集合和结果任务都会被转换为结构化参数，最终交给后端对应接口。前端组件因此起到了参数标准化作用，使用户输入在进入遥感算法前已经具备较清晰的字段含义。"},
        ],
        "4.6.4\u3000结果与下载交互": [
            {"kind": "p", "text": "结果与下载交互强调可追踪和可复查。结果中心按照 processed、composite、mask、metadata 等类别展示产物，用户不需要直接进入磁盘目录识别文件名；下载页面按照场景、资产和任务状态组织信息，用户能够在一个位置看到检索结果、资产选择和服务端下载进度。"},
            {"kind": "p", "text": "这种交互方式也服务于论文测试。功能测试可以围绕页面状态、任务编号、输出目录和清单文件逐项确认，而不必只依赖控制台日志判断系统是否运行成功。对于毕业设计而言，可视化反馈能够更清楚地证明“预处理系统”不仅有算法函数，也具备完整用户操作链。"},
        ],
        "5.1\u3000测试环境与测试方法": [
            {"kind": "p", "text": "本系统测试以功能验证和流程联调为主，不把尚未开展的性能压测或精度对比写入结论。测试环境需要同时满足 Python 后端、Node.js 前端、GDAL 栅格处理库和浏览器运行条件。后端重点检查接口是否能够正确接收参数、创建任务和返回状态；前端重点检查页面参数是否能够正确组织为请求；输出结果重点检查是否生成处理波段、合成图、指数图、清单文件和可预览产物。"},
            {"kind": "p", "text": "测试方法采用“模块测试 + 流程测试 + 结果检查”的组合方式。模块测试用于验证单景、批量、下载和结果中心各自功能；流程测试用于验证从数据输入到结果查看的连续性；结果检查则通过文件目录、任务清单、预览图和页面状态判断功能是否真正落地。对于需要外部网络的 STAC 检索和下载功能，论文只描述系统支持的交互和任务管理逻辑，不编造不可复现的下载速度或检索成功率。"},
        ],
        "5.2\u3000系统功能测试": [
            {"kind": "caption", "text": f"表5.4{IDEOGRAPHIC_SPACE}系统功能测试场景扩展表"},
            {"kind": "table", "headers": ["测试模块", "测试内容", "通过依据"], "rows": [
                ["单景预处理", "上传/选择波段、设置产品级别、提交异步任务", "返回 job_id，状态接口可查询，输出目录生成产物"],
                ["L1/L2 链路", "分别选择 L1 与 L2 产品输入", "L1 执行定标校正，L2 进入缩放换算链路"],
                ["质量控制", "提供 QA_PIXEL 与 QA_RADSAT", "生成掩膜摘要，有效像元比例可追踪"],
                ["批量处理", "构建节点流程并提交", "通过拓扑校验，任务进入队列并更新状态"],
                ["影像下载", "配置 AOI、时间范围和资产选择", "形成检索请求和服务端下载任务"],
                ["结果中心", "扫描当前与历史输出目录", "按类别展示产物，支持预览与下载"],
            ]},
            {"kind": "p", "text": "从功能测试结果看，系统的核心流程可以分为两类：一类是本地数据处理流程，主要依赖本机文件、GDAL 和处理器逻辑；另一类是在线检索下载流程，依赖 STAC 服务、网络状态和认证配置。论文评价时应将两类流程区分说明，避免把网络条件造成的不确定性归因于核心预处理算法。"},
        ],
        "5.3\u3000典型处理结果分析": [
            {"kind": "p", "text": "典型结果分析主要关注结果链是否完整，而不是声称算法精度优于已有方法。处理后波段用于验证基础预处理是否完成，真彩色和假彩色结果用于验证波段组合和显示拉伸是否合理，NDVI、NDWI、NDBI 等指数用于验证公式计算和空间参考继承是否正常，云掩膜结果用于说明异常像元识别能力，多景镶嵌结果用于说明批量流程和同名波段整合能力。"},
            {"kind": "p", "text": "在结果解释上，本文采用定性分析口径：如果输出图像能够保持地物轮廓、空间范围和波段关系一致，并且任务清单能准确记录产物类别，则说明系统已经具备教学实验和本地处理场景所需的基本可用性。后续若要进行更严格的定量评价，可进一步引入标准样区、参考产品和多时相对比实验。"},
        ],
        "5.4\u3000系统运行效果评价": [
            {"kind": "p", "text": "从运行效果看，系统已经形成较完整的用户操作闭环。用户可以从下载页检索或准备影像，也可以直接在单景页面选择本地波段；任务提交后前端能够查看进度，处理完成后结果中心能够汇总输出文件，预览页能够进一步检查栅格或图像结果。该闭环说明系统不只是把命令行脚本搬到后端，而是围绕遥感预处理流程补齐了交互、任务和结果管理环节。"},
            {"kind": "p", "text": "系统运行评价仍需保持克制。由于本次测试主要针对功能流程和典型样例，不具备覆盖不同季节、不同地貌和大规模场景的统计条件，因此本文不写处理速度排名、精度提升百分比或平台级并发能力。更合理的结论是：系统在本地实验环境下能够支撑 Landsat 8/9 影像预处理演示和教学型批量流程，为后续扩展传感器、增加持久化队列和开展定量对比实验提供了基础。"},
        ],
        "6.1\u3000主要工作总结": [
            {"kind": "p", "text": "从工程实现角度看，本文完成的工作不仅包括若干遥感算法函数，还包括围绕这些函数构建的在线工作流。系统通过 FastAPI 将处理能力封装为接口，通过 Vue 3 将复杂参数组织为可操作页面，通过批量流程画布把多景处理抽象为节点图，通过结果资产中心把分散输出重新聚合为可下载、可预览的任务资产。"},
        ],
        "6.2\u3000系统特点与不足": [
            {"kind": "p", "text": "系统的特点主要体现在三个方面：第一，处理链路较完整，覆盖数据输入、预处理、合成指数、批量执行和结果归档；第二，技术边界清晰，前端负责交互和状态展示，后端负责任务与文件处理，算法模块负责栅格计算；第三，工程保护较多，包括路径白名单、6S 失败回退、下载重试、任务清单和结果分类等机制。"},
            {"kind": "p", "text": "系统不足也需要客观说明。当前主预处理链仍以 Landsat 8/9 为核心，Sentinel-2 更多体现为下载与部分处理入口，尚未形成与 Landsat 同等完整的多源处理链；批量任务状态主要保存在内存结构中，跨重启恢复能力有限；测试材料以功能验证为主，尚未形成大规模数据集上的性能统计和定量精度对比。"},
        ],
        "6.3\u3000后续改进方向": [
            {"kind": "p", "text": "后续工作可从工程化和遥感应用两个层面继续推进。工程化方面，可引入轻量数据库或任务日志持久化机制，增强长时间任务的恢复能力；可进一步完善自动化测试、部署脚本和异常日志分析，使系统更适合持续维护。遥感应用方面，可扩展 Sentinel-2、Landsat 7 等产品的主处理链，增加更多指数、分类辅助和统计图表，使平台从预处理工具逐步发展为更完整的本地遥感分析工作台。"},
        ],
    }
    for heading, blocks in enrichments.items():
        paragraph = find_heading(doc, heading)
        if paragraph is not None:
            add_blocks_after(paragraph, blocks)


def remove_bracket_placeholders(doc: Document):
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if re.match(r"^\[\s*图\d+\.\d+.+\]\s*$", text):
            parent = paragraph._element.getparent()
            parent.remove(paragraph._element)


def insert_picture_before(paragraph: Paragraph, image_path: Path, width_cm: float = 14.2):
    if not image_path.exists():
        return
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    pic_paragraph = Paragraph(new_p, paragraph._parent)
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_paragraph.paragraph_format.first_line_indent = Pt(0)
    run = pic_paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def previous_has_picture(doc: Document, index: int) -> bool:
    if index <= 0:
        return False
    prev = doc.paragraphs[index - 1]
    return bool(prev._p.xpath(".//w:drawing"))


def insert_images(doc: Document, assets: dict[str, Path]):
    caption_map = {
        "图1.1": "图1.1",
        "图2.2": "图2.2",
        "图3.1": "图3.1",
        "图4.1": "图4.1",
        "图4.2": "图4.2",
        "图4.3": "图4.3",
        "图4.4": "图4.4",
        "图4.5": "图4.5",
        "图4.6": "图4.6",
        "图4.7": "图4.7",
        "图4.8": "图4.8",
        "图4.9": "图4.9",
        "图5.1": "图5.1",
        "图5.2": "图5.2",
        "图5.3": "图5.3",
        "图5.4": "图5.4",
    }
    for index, paragraph in list(enumerate(doc.paragraphs)):
        text = paragraph.text.strip()
        for prefix, asset_key in caption_map.items():
            if text.startswith(prefix + IDEOGRAPHIC_SPACE) and asset_key in assets and not previous_has_picture(doc, index):
                insert_picture_before(paragraph, assets[asset_key])
                break


def normalize_caption_spacing(doc: Document):
    patterns = [
        (re.compile(r"^(图\d+\.\d+)[\s\u3000]*(.+)$"), "图"),
        (re.compile(r"^(表\d+\.\d+)[\s\u3000]*(.+)$"), "表"),
        (re.compile(r"^(代码清单\d+\.\d+)[\s\u3000]*(.+)$"), "代码"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for pattern, _ in patterns:
            match = pattern.match(text)
            if match:
                new_text = f"{match.group(1)}{IDEOGRAPHIC_SPACE}{match.group(2)}"
                if new_text != text:
                    clear_and_set(paragraph, new_text, cn="黑体" if text.startswith(("图", "表", "代码清单")) else "宋体", size=11)
                    caption_format(paragraph)
                break


def _set_style_font(style, *, cn: str, en: str = "Times New Roman", size: float = 12, bold: bool = False):
    style.font.name = en
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.underline = False
    r_pr = style._element.get_or_add_rPr()
    old = r_pr.find(qn("w:rFonts"))
    if old is not None:
        r_pr.remove(old)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def _set_paragraph_runs_font(paragraph: Paragraph, *, cn: str, en: str = "Times New Roman", size: float = 12, bold: bool = False):
    for run_element in paragraph._p.xpath(".//w:r"):
        r_pr = run_element.get_or_add_rPr()
        r_style = r_pr.find(qn("w:rStyle"))
        if r_style is not None:
            r_pr.remove(r_style)

        old = r_pr.find(qn("w:rFonts"))
        if old is not None:
            r_pr.remove(old)
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:eastAsia"), cn)
        r_fonts.set(qn("w:ascii"), en)
        r_fonts.set(qn("w:hAnsi"), en)
        r_pr.insert(0, r_fonts)

        sz = r_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            r_pr.append(sz)
        sz.set(qn("w:val"), str(int(size * 2)))

        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(int(size * 2)))

        b = r_pr.find(qn("w:b"))
        if bold:
            if b is None:
                b = OxmlElement("w:b")
                r_pr.append(b)
            b.set(qn("w:val"), "1")
        elif b is not None:
            r_pr.remove(b)

        color = r_pr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)
        color.set(qn("w:val"), "000000")

        underline = r_pr.find(qn("w:u"))
        if underline is None:
            underline = OxmlElement("w:u")
            r_pr.append(underline)
        underline.set(qn("w:val"), "none")


def normalize_toc_style(doc: Document):
    toc_specs = {
        "toc 1": {"cn": "黑体", "bold": True, "before": 6, "after": 0},
        "toc 2": {"cn": "宋体", "bold": False, "before": 0, "after": 0},
        "toc 3": {"cn": "宋体", "bold": False, "before": 0, "after": 0},
    }
    for style_name, spec in toc_specs.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        pf = style.paragraph_format
        pf.space_before = Pt(spec["before"])
        pf.space_after = Pt(spec["after"])
        pf.line_spacing = Pt(20)
        _set_style_font(style, cn=spec["cn"], size=12, bold=spec["bold"])

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "目　　录":
            clear_and_set(paragraph, text, cn="黑体", size=15)
            pf = paragraph.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.space_before = Pt(40)
            pf.space_after = Pt(20)
            pf.line_spacing = Pt(20)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        style_name = paragraph.style.name
        if style_name not in toc_specs:
            continue
        spec = toc_specs[style_name]
        pf = paragraph.paragraph_format
        pf.space_before = Pt(spec["before"])
        pf.space_after = Pt(spec["after"])
        pf.line_spacing = Pt(20)
        _set_paragraph_runs_font(paragraph, cn=spec["cn"], size=12, bold=spec["bold"])


def renumber_code_lists_for_reading_order(doc: Document):
    replacements = [
        ("代码清单 4.8", "__AOI_CODE_REF__"),
        ("代码清单4.8", "__AOI_CODE_CAP__"),
        ("代码清单 4.9", "代码清单 4.8"),
        ("代码清单4.9", "代码清单4.8"),
        ("代码清单 4.10", "代码清单 4.9"),
        ("代码清单4.10", "代码清单4.9"),
        ("__AOI_CODE_REF__", "代码清单 4.10"),
        ("__AOI_CODE_CAP__", "代码清单4.10"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != text:
            is_caption = new_text.strip().startswith("代码清单")
            clear_and_set(paragraph, new_text, cn="黑体" if is_caption else "宋体", size=11 if is_caption else 12)
            if is_caption:
                caption_format(paragraph)
            else:
                body_format(paragraph)


def write_note(status: str, counts: dict[str, int]):
    note = OUTPUT.with_name("基于Web的Landsat8遥感影像在线预处理系统-内容丰富版-交付说明.txt")
    lines = [
        "内容丰富版交付说明",
        "",
        f"底稿：{SOURCE}",
        f"输出：{OUTPUT}",
        f"截图状态：{status}",
        "",
        "章节计数字符：",
    ]
    for name, count in counts.items():
        lines.append(f"- {name}: {count}")
    note.write_text("\n".join(lines), encoding="utf-8")


def count_chapters(doc: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    current = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.startswith("Heading 1"):
            current = text
            counts.setdefault(current, 0)
        if current:
            counts[current] += len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9]+", text))
    return counts


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = generate_static_assets()
    screenshot_status = capture_frontend_screenshots(assets)

    doc = Document(SOURCE)
    insert_section_enrichment(doc)
    remove_bracket_placeholders(doc)
    renumber_code_lists_for_reading_order(doc)
    insert_images(doc, assets)
    normalize_caption_spacing(doc)
    normalize_toc_style(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    counts = count_chapters(Document(OUTPUT))
    write_note(screenshot_status, counts)
    print(f"output={OUTPUT}")
    print(f"screenshot_status={screenshot_status}")
    for name, count in counts.items():
        print(f"{name}\t{count}")


if __name__ == "__main__":
    main()
