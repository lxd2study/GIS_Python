from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-内容丰富版.docx"

EXTRA_PARAGRAPHS = {
    "2.2.4　Py6S 大气校正模型接口": [
        "在项目实现中，Py6S 主要服务于 L1 产品的大气校正链路。系统通过元数据和用户配置组织太阳高度角、传感器类型、成像时间和大气参数，再把这些参数传入 6S 模型接口，获得用于反射率校正的计算结果。与直接在正文中手工推导辐射传输模型相比，项目采用 Py6S 更符合工程实现特点，能够把复杂模型封装为可调用模块。",
        "同时，系统没有把 6S 设定为唯一可用路径。考虑到本地环境可能缺少模型依赖、参数不完整或执行失败，后端保留 DOS 方法作为回退方案，并在任务状态中记录处理分支。该设计使论文中的大气校正说明更贴近真实项目：系统重点解决的是在线预处理流程的可运行性和可解释性，而不是声称重新提出新的大气校正算法。",
    ],
    "2.2.5　OpenLayers 与 Vue Flow 交互组件": [
        "OpenLayers 在系统中承担空间交互入口的作用。用户在 AOI 地图组件中绘制矩形或查看矢量范围时，前端会把 Web Mercator 坐标转换为 WGS84 经纬度 bbox，再将该范围写入裁剪参数。这样可以减少手工输入经纬度范围造成的顺序错误和精度问题，也让裁剪范围在提交任务前具有可视化反馈。",
        "Vue Flow 则用于批量处理模块的流程表达。输入数据、辐射定标、大气校正、裁剪、镶嵌、合成和输出等环节被抽象为节点，节点之间的连线表示处理顺序和数据流向。后端 GraphExecutor 会对前端图结构重新做可达性分析和拓扑排序，因此 Vue Flow 不只是界面装饰，而是批量任务配置、校验和执行之间的桥梁。",
    ],
    "2.2.6　STAC 影像检索规范": [
        "STAC 规范用于描述遥感场景、空间范围、时间范围、云量、平台和资产链接等信息，适合将在线遥感数据检索过程标准化。本文系统在影像下载模块中根据传感器和产品级别选择不同集合，例如 Landsat Collection 2 Level-2、Landsat Collection 2 Level-1 和 Sentinel-2 Level-2A，并把检索结果整理为前端可选择的场景列表。",
        "在工程实现上，STAC 检索还需要与认证、资产签名和下载归档结合。对于无需认证的公开资产，系统可以直接生成下载任务；对于需要认证或签名的资产，系统会在服务层区分处理方式，并通过下载任务列表反馈进度和错误信息。这样既保留了在线取数能力，也避免把外部数据平台的不确定性误写成系统内部算法问题。",
    ],
}


def set_run_font(run, cn: str = "宋体", en: str = "Times New Roman", size: float = 12) -> None:
    run.font.name = en
    run.font.size = Pt(size)
    r_pr = run._r.get_or_add_rPr()
    old = r_pr.find(qn("w:rFonts"))
    if old is not None:
        r_pr.remove(old)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def body_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(20)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def insert_after(anchor, text: str):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = anchor.__class__(new_p, anchor._parent)
    run = paragraph.add_run(text)
    set_run_font(run)
    body_format(paragraph)
    return paragraph


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def find_heading(doc: Document, heading: str):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == heading:
            return paragraph
    return None


def main() -> None:
    doc = Document(DOCX)
    for heading, paragraphs in EXTRA_PARAGRAPHS.items():
        anchor = find_heading(doc, heading)
        if anchor is None:
            continue
        section_paragraphs = []
        cursor = anchor
        while cursor._p.getnext() is not None:
            nxt = cursor._p.getnext()
            if nxt.tag != qn("w:p"):
                break
            candidate = cursor.__class__(nxt, cursor._parent)
            if candidate.style.name.startswith("Heading"):
                break
            if candidate.text.strip():
                section_paragraphs.append(candidate)
            cursor = candidate

        for paragraph in list(section_paragraphs):
            if paragraph.text.strip() in paragraphs:
                remove_paragraph(paragraph)

        anchor_for_insert = next((p for p in section_paragraphs if p.text.strip() not in paragraphs), anchor)
        cursor = anchor_for_insert
        for text in paragraphs:
            cursor = insert_after(cursor, text)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
