from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-内容丰富版.docx"
IDEOGRAPHIC_SPACE = "\u3000"


def count_chapters(doc: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    current = ""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.startswith("Heading 1"):
            current = text
            counts.setdefault(current, 0)
        if current:
            counts[current] += len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9]+", text))
    return counts


def main() -> int:
    doc = Document(DOCX)
    texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    counts = count_chapters(doc)
    longest = max(counts.items(), key=lambda item: item[1])

    heading_bad: list[str] = []
    caption_bad: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name.startswith("Heading") and re.match(r"^(第\d+章|\d+(?:\.\d+)+) ", text):
            heading_bad.append(text)
        if re.match(r"^(图|表|代码清单)\d+\.\d+ ", text):
            caption_bad.append(text)

    placeholders = [text for text in texts if any(term in text for term in ["待补充", "占位图"])]
    figure_captions = [text for text in texts if re.match(rf"^图\d+\.\d+{IDEOGRAPHIC_SPACE}", text)]
    table_captions = [text for text in texts if re.match(rf"^表\d+\.\d+{IDEOGRAPHIC_SPACE}", text)]
    code_captions = [text for text in texts if re.match(rf"^代码清单\d+\.\d+{IDEOGRAPHIC_SPACE}", text)]
    drawings = len(doc.element.xpath(".//w:drawing"))

    checks = {
        "docx_exists": DOCX.exists(),
        "chapter4_longest": longest[0].startswith("第4章"),
        "code_caption_count_10_to_12": 10 <= len(code_captions) <= 12,
        "no_heading_halfwidth_space": not heading_bad,
        "no_caption_halfwidth_space": not caption_bad,
        "no_placeholder_terms": not placeholders,
        "has_images": drawings >= 10,
    }

    print(f"file={DOCX}")
    print(f"size={DOCX.stat().st_size}")
    print(f"chapter_counts={counts}")
    print(f"longest={longest[0]}:{longest[1]}")
    print(f"total_main_count={sum(value for key, value in counts.items() if key.startswith('第'))}")
    print(f"figure_captions={len(figure_captions)}")
    print(f"table_captions={len(table_captions)}")
    print(f"code_captions={len(code_captions)}")
    for caption in code_captions:
        print(f"  {caption}")
    print(f"drawings={drawings}")
    print(f"heading_bad={heading_bad[:10]}")
    print(f"caption_bad={caption_bad[:10]}")
    print(f"placeholders={placeholders[:10]}")
    print(f"checks={checks}")
    return 0 if all(checks.values()) else 1


if __name__ == "__main__":
    raise SystemExit(main())
