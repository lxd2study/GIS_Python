from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-内容丰富版.docx"


def set_run_font(run, *, cn: str, en: str = "Times New Roman", size: float = 12, bold: bool = False) -> None:
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.underline = False
    r_pr = run._r.get_or_add_rPr()
    old = r_pr.find(qn("w:rFonts"))
    if old is not None:
        r_pr.remove(old)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def set_style_font(style, *, cn: str, en: str = "Times New Roman", size: float = 12, bold: bool = False) -> None:
    style.font.name = en
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.underline = False
    r_pr = style._element.get_or_add_rPr()
    old = r_pr.find(qn("w:rFonts"))
    if old is not None:
        r_pr.remove(old)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def set_xml_runs_font(paragraph, *, cn: str, en: str = "Times New Roman", size: float = 12, bold: bool = False) -> None:
    for run_element in paragraph._p.xpath(".//w:r"):
        r_pr = run_element.get_or_add_rPr()
        r_style = r_pr.find(qn("w:rStyle"))
        if r_style is not None:
            r_pr.remove(r_style)

        old = r_pr.find(qn("w:rFonts"))
        if old is not None:
            r_pr.remove(old)
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:eastAsia"), cn)
        r_fonts.set(qn("w:ascii"), en)
        r_fonts.set(qn("w:hAnsi"), en)
        r_pr.insert(0, r_fonts)

        for tag in ("w:sz", "w:szCs"):
            node = r_pr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                r_pr.append(node)
            node.set(qn("w:val"), str(int(size * 2)))

        bold_node = r_pr.find(qn("w:b"))
        if bold:
            if bold_node is None:
                bold_node = OxmlElement("w:b")
                r_pr.append(bold_node)
            bold_node.set(qn("w:val"), "1")
        elif bold_node is not None:
            r_pr.remove(bold_node)

        color = r_pr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            r_pr.append(color)
        color.set(qn("w:val"), "000000")

        underline = r_pr.find(qn("w:u"))
        if underline is None:
            underline = OxmlElement("w:u")
            r_pr.append(underline)
        underline.set(qn("w:val"), "none")


def normalize_toc_style(doc: Document) -> None:
    toc_specs = {
        "toc 1": {"cn": "黑体", "bold": True, "before": 6, "after": 0},
        "toc 2": {"cn": "宋体", "bold": False, "before": 0, "after": 0},
        "toc 3": {"cn": "宋体", "bold": False, "before": 0, "after": 0},
    }
    for style_name, spec in toc_specs.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        pf = style.paragraph_format
        pf.space_before = Pt(spec["before"])
        pf.space_after = Pt(spec["after"])
        pf.line_spacing = Pt(20)
        set_style_font(style, cn=spec["cn"], size=12, bold=spec["bold"])

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "目　　录":
            paragraph.clear()
            run = paragraph.add_run(text)
            set_run_font(run, cn="黑体", en="Times New Roman", size=15, bold=False)
            pf = paragraph.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.space_before = Pt(40)
            pf.space_after = Pt(20)
            pf.line_spacing = Pt(20)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        spec = toc_specs.get(paragraph.style.name)
        if not spec:
            continue
        pf = paragraph.paragraph_format
        pf.space_before = Pt(spec["before"])
        pf.space_after = Pt(spec["after"])
        pf.line_spacing = Pt(20)
        set_xml_runs_font(paragraph, cn=spec["cn"], size=12, bold=spec["bold"])


def main() -> None:
    doc = Document(DOCX)
    normalize_toc_style(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
