from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "thesis-prep" / "基于Web的Landsat8遥感影像在线预处理系统.docx"
OUTPUT = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-标题技术代码优化版.docx"
IDEOGRAPHIC_SPACE = "\u3000"


TITLE_MAP = {
    "遥感影像预处理相关理论": "遥感影像预处理理论基础",
    "区域裁剪与结果表达": "区域裁剪与专题结果表达",
    "系统开发相关技术": "系统关键开发技术",
    "FastAPI": "FastAPI 后端接口框架",
    "Vue 3 与 Vite": "Vue 3 与 Vite 前端构建技术",
    "GDAL 与 NumPy": "GDAL 与 NumPy 栅格处理技术",
    "Py6S": "Py6S 大气校正模型接口",
    "OpenLayers 与 Vue Flow": "OpenLayers 与 Vue Flow 交互组件",
    "STAC 检索": "STAC 影像检索规范",
    "系统需求分析": "系统功能与非功能需求分析",
    "可行性分析": "系统可行性分析",
    "系统建设目标与设计原则": "系统建设目标及设计原则",
    "后端结构设计": "后端分层结构设计",
    "前端工作台结构": "前端工作台结构设计",
    "功能模块设计": "系统功能模块设计",
    "主要后端接口说明表": "核心后端接口说明表",
    "优先级队列与 worker 执行机制": "优先级队列与 Worker 执行机制",
    "STAC 检索与集合配置": "STAC 检索与数据集合配置",
    "补充处理逻辑图": "核心处理逻辑图设计",
    "功能测试": "系统功能测试",
    "系统运行评价": "系统运行效果评价",
}


TECH_SECTION = [
    {
        "kind": "heading",
        "text": f"2.2.7{IDEOGRAPHIC_SPACE}项目关键技术应用",
        "style": "Heading 3",
    },
    {
        "kind": "body",
        "text": "结合项目源码可知，本系统的关键技术并不是简单罗列开发框架，而是分别落实在接口服务、栅格处理、批量编排、影像检索和前端交互等模块中。后端在 `app.py` 中统一创建 FastAPI 应用，并注入进度管理、文件管理、批量任务和影像下载服务；核心处理器在 `processor.py` 中根据产品级别区分 L1 与 L2 链路；批量处理由前端 Vue Flow 生成图结构，再由 `GraphExecutor` 进行可达性分析和拓扑排序；影像下载服务在 `landsat_download.py` 中维护不同传感器和产品级别的 STAC 集合配置；前端则通过 OpenLayers 完成 AOI 绘制、范围转换与地图交互。",
    },
    {
        "kind": "body",
        "text": "因此，本文后续实现章节重点围绕“关键技术如何在源码中落地”展开说明。通过将核心代码片段嵌入第 4 章，可以使系统设计、功能描述和项目实现之间形成更直接的对应关系，也便于答辩时说明各模块并非停留在概念设计层面，而是已经完成了可运行的工程实现。",
    },
]


INSERTIONS = {
    "在接口设计上，系统提供异步预处理接口接收上述参数。": [
        {
            "kind": "body",
            "text": "代码清单 4.1 展示了单景异步预处理接口的核心组织方式。接口接收波段、MTL、QA、裁剪范围、合成类型和产品级别等参数，创建任务编号后将文件准备和实际预处理拆分为两个阶段，从而支撑前端轮询进度与后台执行任务。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.1{IDEOGRAPHIC_SPACE}异步预处理接口创建任务片段",
        },
        {
            "kind": "code",
            "text": """@app.post("/preprocess_landsat8_async")
async def preprocess_landsat8_async(..., product_level: str = Form("L1")) -> Dict:
    job_id = str(uuid.uuid4())
    progress_manager.init_progress(job_id)
    temp_dir = file_manager.create_temp_dir(prefix=f"landsat8_{job_id}_")
    preprocess_inputs = await _prepare_async_preprocess_inputs(...)
    _launch_async_preprocess(
        job_id=job_id,
        band_paths=preprocess_inputs["band_paths"],
        product_level=product_level,
        cleanup_temp_dir=temp_dir,
    )
    return {"job_id": job_id, "status": "processing"}""",
        },
    ],
    "对于 L2 产品，系统不再重复进行辐射定标和大气校正": [
        {
            "kind": "body",
            "text": "代码清单 4.2 对应 L2 表面反射率缩放逻辑。该片段先保留有效像元，再应用 Landsat Collection 2 Level-2 产品的比例系数与偏移量，并对异常值进行空值化和范围约束，保证后续合成和指数计算使用的是稳定的浮点反射率数组。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.2{IDEOGRAPHIC_SPACE}L2 表面反射率缩放处理片段",
        },
        {
            "kind": "code",
            "text": """reflectance = np.asarray(raw_values, dtype=np.float32)
valid_mask = np.isfinite(reflectance) & (reflectance > 0)
np.multiply(reflectance, LANDSAT_L2_SR_SCALE, out=reflectance)
reflectance += LANDSAT_L2_SR_OFFSET
reflectance[~valid_mask] = np.nan
if np.any(valid_mask):
    reflectance[valid_mask] = np.clip(reflectance[valid_mask], -0.2, 1.6)""",
        },
    ],
    "若用户选择 DOS，系统直接对 TOA 反射率执行暗目标扣除": [
        {
            "kind": "body",
            "text": "代码清单 4.3 展示了 6S 与 DOS 的工程化衔接。系统优先按用户选择执行 6S 模型；当模型运行环境或参数条件导致失败时，自动回退到 DOS 方法，避免单个波段校正失败造成整个任务中断。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.3{IDEOGRAPHIC_SPACE}6S 失败回退到 DOS 的处理片段",
        },
        {
            "kind": "code",
            "text": """try:
    corrected = self.sixs_atmospheric_correction(reflectance, band_name)
    return corrected, "6S"
except Exception as exc:
    logger.warning("波段 %s: 6S大气校正失败,回退到DOS方法", band_name)
    corrected = dark_object_subtraction(reflectance)
    return corrected, "DOS(6S失败回退)" """,
        },
    ],
    "系统支持真彩色、假彩色、农业监测、城市研究、短波红外等预设合成": [
        {
            "kind": "body",
            "text": "对于自定义指数，系统没有直接执行用户输入的字符串，而是通过 Python AST 限制可用语法和函数，避免任意代码执行风险。代码清单 4.4 体现了公式解析、波段变量提取和数组表达式求值的基本流程。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.4{IDEOGRAPHIC_SPACE}自定义指数安全解析片段",
        },
        {
            "kind": "code",
            "text": """expr = ast.parse(formula, mode="eval")
for node in ast.walk(expr):
    if isinstance(node, ast.Name):
        band_name = node.id.upper()
        if BAND_NAME_RE.match(band_name):
            band_names.add(band_name)
        elif node.id not in ALLOWED_FUNCTIONS:
            raise Exception(f"Unsupported symbol in formula: {node.id}")
with np.errstate(divide="ignore", invalid="ignore"):
    result = _eval_formula(expr, band_arrays)""",
        },
    ],
    "只有当流程图通过校验后，系统才会为单场景任务或镶嵌任务生成对应的 `BatchJobConfig` 列表。": [
        {
            "kind": "body",
            "text": "代码清单 4.5 给出了图结构拓扑排序的核心过程。该算法根据连线统计每个节点的入度，从入度为零的节点开始出队，并在遍历过程中逐步减少后继节点入度，从而得到满足依赖关系的执行顺序。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.5{IDEOGRAPHIC_SPACE}批量流程拓扑排序片段",
        },
        {
            "kind": "code",
            "text": """for edge in edges:
    if edge["source"] in node_ids and edge["target"] in node_ids:
        adj[edge["source"]].append(edge["target"])
        in_degree[edge["target"]] += 1
queue = deque([node["id"] for node in nodes if in_degree[node["id"]] == 0])
while queue:
    node_id = queue.popleft()
    result.append(node_id)
    for nxt in adj[node_id]:
        in_degree[nxt] -= 1""",
        },
    ],
    "虽然系统当前没有引入数据库或分布式调度组件": [
        {
            "kind": "body",
            "text": "代码清单 4.6 展示了批量任务管理器的队列和 Worker 初始化方式。系统按照高、中、低优先级分别维护队列，并在启动时创建后台线程持续消费任务，满足本地批处理场景下的排队执行需求。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.6{IDEOGRAPHIC_SPACE}批量任务优先级队列与 Worker 启动片段",
        },
        {
            "kind": "code",
            "text": """self.job_queues = {
    TaskPriority.HIGH: queue.PriorityQueue(),
    TaskPriority.MEDIUM: queue.PriorityQueue(),
    TaskPriority.LOW: queue.PriorityQueue(),
}
for i in range(self.max_workers):
    worker = threading.Thread(target=self._worker_loop, args=(i,), daemon=True)
    worker.start()
    self.workers.append(worker)""",
        },
    ],
    "包括 Landsat L1、Landsat L2、Landsat 7 L1/L2 以及 Sentinel-2 L2A 等选项": [
        {
            "kind": "body",
            "text": "代码清单 4.7 体现了检索模块对不同产品集合的配置方式。L2 产品使用 Planetary Computer STAC 服务并进行签名，L1 产品使用 USGS STAC 服务并标记认证要求，前端据此展示产品说明并选择对应下载流程。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.7{IDEOGRAPHIC_SPACE}STAC 数据集合配置片段",
        },
        {
            "kind": "code",
            "text": """"L2": {
    "url": "https://planetarycomputer.microsoft.com/api/stac/v1",
    "collection": "landsat-c2-l2",
    "sign": True,
    "auth_required": False,
},
"L1": {
    "url": "https://landsatlook.usgs.gov/stac-server",
    "collection": "landsat-c2l1",
    "sign": False,
    "auth_required": True,
}""",
        },
    ],
    "AOI 交互主要由 OpenLayers 完成。": [
        {
            "kind": "body",
            "text": "代码清单 4.8 展示了前端 AOI 框选的核心逻辑。前端通过 OpenLayers 的 Draw 交互创建矩形，绘制结束后把 Web Mercator 坐标转换为 WGS84 经纬度范围，再通过组件事件传递给任务参数。",
        },
        {
            "kind": "caption",
            "text": f"代码清单4.8{IDEOGRAPHIC_SPACE}AOI 矩形绘制与坐标转换片段",
        },
        {
            "kind": "code",
            "text": """drawInteraction = new Draw({
  source: selectionSource,
  type: "Circle",
  geometryFunction: createBox(),
})
drawInteraction.on("drawend", (event) => {
  const bbox = transformExtent(
    event.feature.getGeometry().getExtent(),
    "EPSG:3857",
    "EPSG:4326",
  ).map((value) => Number(value.toFixed(6)))
  emit("update:modelValue", bbox.join(","))
})""",
        },
    ],
}


def set_run_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, color=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

    r_pr = run._r.get_or_add_rPr()
    existing = r_pr.find(qn("w:rFonts"))
    if existing is not None:
        r_pr.remove(existing)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), en)
    r_fonts.set(qn("w:hAnsi"), en)
    r_pr.insert(0, r_fonts)


def clear_and_set_text(paragraph: Paragraph, text: str, *, cn="宋体", en="Times New Roman", size=12, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)


def normalized_title(title: str) -> str:
    compact = re.sub(r"[\s\u3000]+", " ", title).strip()
    return TITLE_MAP.get(compact, compact)


def format_numbered_title(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return stripped

    special_titles = {
        "摘 要": f"摘{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}要",
        "摘  要": f"摘{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}要",
        "摘要": f"摘{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}要",
        "目 录": f"目{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}录",
        "目  录": f"目{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}录",
        "目录": f"目{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}录",
        "致谢": f"致{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}谢",
        "致 谢": f"致{IDEOGRAPHIC_SPACE}{IDEOGRAPHIC_SPACE}谢",
    }
    if stripped in special_titles:
        return special_titles[stripped]

    chapter = re.match(r"^(第\s*(\d+)\s*章)[\s\u3000]*(.+)$", stripped)
    if chapter:
        return f"第{chapter.group(2)}章{IDEOGRAPHIC_SPACE}{normalized_title(chapter.group(3))}"

    section = re.match(r"^(\d+(?:\.\d+)+)[\s\u3000]*(.+)$", stripped)
    if section:
        return f"{section.group(1)}{IDEOGRAPHIC_SPACE}{normalized_title(section.group(2))}"

    appendix = re.match(r"^(附录[A-ZＡ-Ｚ])[\s\u3000]*(.+)$", stripped)
    if appendix:
        return f"{appendix.group(1)}{IDEOGRAPHIC_SPACE}{normalized_title(appendix.group(2))}"

    caption = re.match(r"^((?:图|表)\d+(?:[.-]\d+)+)[\s\u3000]*(.+)$", stripped)
    if caption:
        return f"{caption.group(1)}{IDEOGRAPHIC_SPACE}{normalized_title(caption.group(2))}"

    return normalized_title(stripped)


def format_toc_text(text: str) -> str:
    if not text.strip():
        return text
    parts = text.rsplit("\t", 1)
    if len(parts) == 2 and parts[1].strip().isdigit():
        return f"{format_numbered_title(parts[0])}\t{parts[1].strip()}"
    match = re.match(r"^(.+?)(\s+)(\d+)$", text.strip())
    if match:
        return f"{format_numbered_title(match.group(1))}\t{match.group(3)}"
    return format_numbered_title(text)


def normalize_titles(doc: Document) -> int:
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text.strip():
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_toc = style_name.lower().startswith("toc") or style_name.lower().startswith("目录")
        is_heading = style_name.startswith("Heading") or style_name.startswith("标题")
        looks_numbered = bool(
            re.match(r"^\s*(第\s*\d+\s*章|\d+(?:\.\d+)+|图\d+(?:[.-]\d+)+|表\d+(?:[.-]\d+)+|附录)", text)
        )
        looks_special = text.strip() in {"摘要", "摘 要", "摘  要", "目 录", "目  录", "目录", "致谢", "致 谢"}

        if not (is_toc or is_heading or looks_numbered or looks_special):
            continue

        new_text = format_toc_text(text) if is_toc else format_numbered_title(text)
        if new_text == text:
            continue

        if is_heading:
            if style_name.endswith("1"):
                clear_and_set_text(paragraph, new_text, cn="黑体", size=15, bold=True)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style_name.endswith("2"):
                clear_and_set_text(paragraph, new_text, cn="黑体", size=14, bold=True)
            else:
                clear_and_set_text(paragraph, new_text, cn="黑体", size=13, bold=True)
        elif is_toc:
            clear_and_set_text(paragraph, new_text, cn="黑体" if "1" in style_name else "宋体", size=12, bold="1" in style_name)
        elif re.match(r"^\s*(图|表)\d+", text):
            clear_and_set_text(paragraph, new_text, cn="黑体", size=11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            clear_and_set_text(paragraph, new_text)

        changed += 1
    return changed


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_body_format(paragraph: Paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(20)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_heading_format(paragraph: Paragraph, level: int):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.line_spacing = Pt(20)
    pf.space_before = Pt({1: 40, 2: 24, 3: 12}.get(level, 12))
    pf.space_after = Pt({1: 20, 2: 6, 3: 6}.get(level, 6))
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size={1: 15, 2: 14, 3: 13}.get(level, 12), bold=True)


def set_caption_format(paragraph: Paragraph):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(16)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size=11)


def shade_paragraph(paragraph: Paragraph, fill: str = "F5F5F5"):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_insert_block(after: Paragraph, block: list[dict]) -> Paragraph:
    cursor = after
    for item in block:
        kind = item["kind"]
        if kind == "heading":
            paragraph = insert_paragraph_after(cursor, item["text"], item.get("style", "Heading 3"))
            level_match = re.search(r"(\d+)$", item.get("style", "Heading 3"))
            set_heading_format(paragraph, int(level_match.group(1)) if level_match else 3)
        elif kind == "body":
            paragraph = insert_paragraph_after(cursor, "")
            clear_and_set_text(paragraph, item["text"], cn="宋体", size=12)
            set_body_format(paragraph)
        elif kind == "caption":
            paragraph = insert_paragraph_after(cursor, "")
            clear_and_set_text(paragraph, item["text"], cn="黑体", size=11)
            set_caption_format(paragraph)
        elif kind == "code":
            paragraph = cursor
            for line in item["text"].splitlines():
                paragraph = insert_paragraph_after(paragraph, "")
                run = paragraph.add_run(line if line else " ")
                set_run_font(run, cn="Courier New", en="Courier New", size=9, color=RGBColor(0, 0, 0))
                pf = paragraph.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Cm(0.55)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = Pt(12)
                shade_paragraph(paragraph)
            cursor = paragraph
            continue
        else:
            raise ValueError(f"Unsupported block kind: {kind}")
        cursor = paragraph
    return cursor


def find_first_paragraph(doc: Document, needle: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def insert_toc_entry(doc: Document):
    if any(
        "2.2.7" in paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style is not None and paragraph.style.name.lower().startswith("toc")
    ):
        return
    anchor = find_first_paragraph(doc, "2.2.6")
    if anchor is None:
        return
    page = "8"
    match = re.search(r"(\d+)\s*$", anchor.text)
    if match:
        page = match.group(1)
    new_para = insert_paragraph_after(anchor, f"2.2.7{IDEOGRAPHIC_SPACE}项目关键技术应用\t{page}", anchor.style.name)
    clear_and_set_text(new_para, new_para.text, cn="宋体", size=12)


def insert_content(doc: Document) -> int:
    inserted = 0

    if not any("项目关键技术应用" in paragraph.text for paragraph in doc.paragraphs):
        anchor = find_first_paragraph(doc, "STAC（Spatiotemporal Asset Catalog）")
        if anchor is not None:
            add_insert_block(anchor, TECH_SECTION)
            inserted += 1
            insert_toc_entry(doc)

    for needle, block in INSERTIONS.items():
        if any(block_item.get("text", "") in paragraph.text for block_item in block if block_item.get("kind") == "caption" for paragraph in doc.paragraphs):
            continue
        anchor = find_first_paragraph(doc, needle)
        if anchor is None:
            continue
        add_insert_block(anchor, block)
        inserted += 1
    return inserted


def main() -> None:
    doc = Document(SOURCE)
    title_changes = normalize_titles(doc)
    content_blocks = insert_content(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"source={SOURCE}")
    print(f"output={OUTPUT}")
    print(f"title_changes={title_changes}")
    print(f"content_blocks={content_blocks}")


if __name__ == "__main__":
    main()
