"""
将毕业设计论文 Markdown 转换为格式规范的 Word (.docx) 文件。

用法:
    python tools/md_to_docx.py docs/毕业设计论文_v1.md docs/毕业设计论文_v1.docx

功能:
- 标题1-4 → Word 标题样式（带自动编号兼容）
- 正文段落（含缩进）
- 有序/无序列表
- 表格（| 分隔符语法）
- 代码块（``` 围栏）
- 占位图注释行（[占位图 X-X]）→ 灰色占位框
- 粗体 **text** 内联渲染
- 公式行（含 $$ 或 \tag）→ 等宽段落保留
- 分隔线 --- → 段落分隔
- 页眉（学校/题目）与页脚（页码）
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────── 辅助函数 ────────────────────────────

def set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=12, bold=False,
             italic=False, color=None):
    """同时设置中英文字体、字号、粗体、斜体、颜色。"""
    run.font.name = name_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # 设置中文字体（东亚字体 rFonts）
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def clear_run_color(run):
    if run.font.color:
        run.font.color.rgb = None


def set_paragraph_format(para, first_indent_chars=2, space_before=0,
                          space_after=6, line_spacing_pt=None,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """设置段落格式：首行缩进、段前/后间距、行距、对齐。"""
    pf = para.paragraph_format
    pf.first_line_indent = Pt(12 * first_indent_chars)  # 12pt ≈ 一个字符宽度
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing_pt:
        pf.line_spacing = Pt(line_spacing_pt)
    pf.alignment = alignment


def add_heading(doc, text, level):
    """添加标题段落（Word 内置 Heading 样式）。"""
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    style = style_map.get(level, 'Heading 4')
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)

    size_map = {1: 15, 2: 14, 3: 13, 4: 12}
    bold_map = {1: True, 2: True, 3: True, 4: True}
    font_map = {1: "黑体", 2: "黑体", 3: "黑体", 4: "黑体"}
    en_font = "Arial" if level == 1 and text == "ABSTRACT" else "Times New Roman"

    set_font(run, name_cn=font_map[level], name_en=en_font,
             size_pt=size_map[level], bold=bold_map[level])

    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt({1: 40, 2: 24, 3: 12, 4: 12}[level])
    pf.space_after = Pt({1: 20, 2: 6, 3: 6, 4: 6}[level])
    pf.line_spacing = Pt(20)
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_body_paragraph(doc, text, first_indent=2):
    """添加正文段落（宋体/Times New Roman 12pt，首行缩进2字符）。"""
    para = doc.add_paragraph()
    _apply_inline(para, text, size_pt=12)
    set_paragraph_format(para, first_indent_chars=first_indent,
                         space_before=0, space_after=0, line_spacing_pt=20)
    return para


def add_special_paragraph(doc, text, section_name="", first_indent=2):
    if section_name == "参考文献" and re.match(r"^\[\d+\]", text):
        para = doc.add_paragraph()
        _apply_inline(para, text, size_pt=10.5)
        pf = para.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(21)
        pf.first_line_indent = Pt(-21)
        pf.space_before = Pt(3)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(16)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para

    if section_name in {"摘 要", "ABSTRACT"}:
        size = 12
        cn_font = "宋体"
        en_font = "Times New Roman"
        para = doc.add_paragraph()
        _apply_inline(para, text, size_pt=size)
        pf = para.paragraph_format
        pf.first_line_indent = Pt(0 if first_indent == 0 else 24)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(20)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            set_font(run, name_cn=cn_font, name_en=en_font, size_pt=size, bold=run.font.bold)
        return para

    return add_body_paragraph(doc, text, first_indent=first_indent)


def add_caption(doc, text, before_pt=6, after_pt=12, caption_type="figure"):
    """图表题注（居中，黑体11pt）。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, name_cn="黑体", name_en="Times New Roman", size_pt=11)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_type == "table":
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    else:
        para.paragraph_format.space_before = Pt(before_pt)
        para.paragraph_format.space_after = Pt(after_pt)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.line_spacing = Pt(12)
    return para


def add_placeholder_box(doc, caption_text):
    """插入灰色占位图框 + 图标注。"""
    # 占位段落（居中，正式稿默认黑色）
    para = doc.add_paragraph()
    run = para.add_run(f"[ {caption_text} ]")
    set_font(run, name_cn="黑体", name_en="Arial", size_pt=11)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.first_line_indent = Pt(0)
    return para


def add_image_with_caption(doc, image_path: Path, caption_text: str):
    """插入图片并添加图题；图片缺失时退回占位框。"""
    if image_path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        try:
            run = para.add_run()
            run.add_picture(str(image_path), width=Cm(14.5))
        except Exception:
            add_placeholder_box(doc, caption_text)
    else:
        add_placeholder_box(doc, caption_text)
    add_caption(doc, caption_text)


def normalize_heading_text(text):
    """清洗章节标题中的说明性后缀，保留正式标题文本。"""
    cleaned = text.strip()
    cleaned = re.sub(r'[（(](说明|占位)[^）)]*[）)]$', '', cleaned).strip()
    if cleaned in {"摘要", "摘 要", "摘  要"}:
        return "摘  要"
    if cleaned.upper() == "ABSTRACT":
        return "ABSTRACT"
    if cleaned.startswith("目 录") or cleaned.startswith("目  录") or cleaned.startswith("目录"):
        return "目  录"
    return cleaned


def extract_table_caption(text):
    """识别表题，兼容表1.1 / 表1-1 形式。"""
    match = re.match(r'^\*{0,2}(表\d+(?:[.-]\d+)+(?:\s+.*)?)\*{0,2}$', text.strip())
    return match.group(1).strip() if match else None


def extract_figure_caption(text):
    """识别 blockquote 形式的占位图题。"""
    cleaned = text.strip()
    if cleaned.startswith('>'):
        cleaned = cleaned.lstrip('> ').strip().strip('*')
    match = re.match(r'^(图\d+(?:[.-]\d+)+(?:\s+.*)?)$', cleaned)
    if not match:
        return None
    caption = match.group(1).strip()
    caption = re.split(r'[（(]占位图', caption, maxsplit=1)[0].strip()
    return caption


def extract_cover_info(lines):
    """从 Markdown 总稿中提取封面信息，未提供时留空。"""
    cover_info = {
        "题目": "",
        "英文题目": "",
        "学生姓名": "",
        "学号": "",
        "指导教师": "",
        "专业名称": "",
        "所在学院": "",
        "完成日期": "",
        "打印日期": "",
    }

    table_key_map = {
        "学生姓名": "学生姓名",
        "学号": "学号",
        "指导教师": "指导教师",
        "专业": "专业名称",
        "专业名称": "专业名称",
        "院（系）": "所在学院",
        "学院": "所在学院",
        "所在学院": "所在学院",
        "完成日期": "完成日期",
        "打印日期": "打印日期",
    }

    for raw_line in lines[:80]:
        line = raw_line.strip()
        if line.startswith("论文题目："):
            cover_info["题目"] = line.split("：", 1)[1].strip()
            continue
        if line.startswith("英文题目："):
            cover_info["英文题目"] = line.split("：", 1)[1].strip()
            continue
        if re.match(r'^\|.*\|.*\|$', line):
            cells = parse_table_line(line)
            if len(cells) >= 2:
                key = cells[0].strip()
                value = cells[1].strip()
                mapped_key = table_key_map.get(key)
                if mapped_key and value not in {"", "内容", "[待填写]"}:
                    cover_info[mapped_key] = value

    return cover_info


def add_code_block(doc, code_text):
    """添加代码块（Courier New 9pt，浅灰背景）。"""
    for line in code_text.split('\n'):
        para = doc.add_paragraph()
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # 浅灰背景
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)


def add_formula_line(doc, text):
    """公式行：等宽字体居中显示。"""
    para = doc.add_paragraph()
    run = para.add_run(text.strip())
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    return para


def add_table(doc, header_row, data_rows):
    """插入带表头的表格（三线表风格）。"""
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(table)

    # 表头
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        cell_text = cell_text.strip().strip('*')
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(cell_text)
        set_font(run, name_cn="黑体", name_en="Times New Roman",
                 size_pt=11, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
        hdr_cells[i].paragraphs[0].paragraph_format.space_after = Pt(3)

    # 数据行
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell_text = cell_text.strip()
            row_cells[col_idx].text = ''
            run = row_cells[col_idx].paragraphs[0].add_run(cell_text)
            set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=11)
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[col_idx].paragraphs[0].paragraph_format.space_before = Pt(3)
            row_cells[col_idx].paragraphs[0].paragraph_format.space_after = Pt(3)

    set_table_border(table, "top", "single", "000000", "12")
    set_table_border(table, "bottom", "single", "000000", "12")
    if len(table.rows) > 1:
        set_row_bottom_border(table.rows[0], "single", "000000", "8")

    # 表格上下加空行
    doc.add_paragraph()
    return table


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)


def set_table_border(table, edge, val, color, size):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_row_bottom_border(row, val, color, size):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), val)
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), color)


def _apply_inline(para, text, size_pt=12, base_bold=False):
    """处理行内粗体 **text** 及普通文本，添加到段落。"""
    # 拆分 **bold** 片段
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, name_cn="宋体", name_en="Times New Roman",
                     size_pt=size_pt, bold=True)
        else:
            if part:
                run = para.add_run(part)
                set_font(run, name_cn="宋体", name_en="Times New Roman",
                         size_pt=size_pt, bold=base_bold)


def add_list_item(doc, text, ordered=False, number=1, level=0):
    """添加列表项。"""
    para = doc.add_paragraph()
    indent_cm = 0.5 + level * 0.5
    if ordered:
        prefix = f"{number}. "
    else:
        prefix = "• "
    run_prefix = para.add_run(prefix)
    set_font(run_prefix, name_cn="宋体", name_en="Times New Roman", size_pt=12)
    _apply_inline(para, text, size_pt=12)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Cm(indent_cm)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    return para


def add_page_number(doc):
    """在页脚中添加居中页码。"""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # 插入域代码 PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'


def setup_page(doc):
    """设置A4页面、页边距。"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = True


def add_header(doc, header_text="河北水利电力学院本科毕业设计"):
    """添加页眉（封面页除外）。"""
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(header_text)
    set_font(run, name_cn="宋体", name_en="Times New Roman", size_pt=10.5)


def add_title_page(doc, title, author, advisor, major, college, date, title_en="", student_id="", print_date=""):
    """生成封面页。"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    set_font(run, name_cn="黑体", name_en="Times New Roman", size_pt=26, bold=True)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(36)
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.line_spacing = Pt(24)

    if title_en:
        p_title_en = doc.add_paragraph()
        run_en = p_title_en.add_run(title_en)
        set_font(run_en, name_cn="Times New Roman", name_en="Times New Roman",
                 size_pt=22, bold=True)
        p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title_en.paragraph_format.space_after = Pt(28)
        p_title_en.paragraph_format.first_line_indent = Pt(0)
        p_title_en.paragraph_format.line_spacing = Pt(20)

    for label, value in [
        ("学生姓名", author),
        ("学号", student_id),
        ("院（系）", college),
        ("专业", major),
        ("指导教师", advisor),
        ("完成日期", date),
        ("打印日期", print_date),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}：")
        font_cn = "宋体" if label in {"完成日期", "打印日期"} else "仿宋"
        if label == "学号":
            font_cn = "Times New Roman"
        set_font(run_label, name_cn=font_cn, name_en="Times New Roman",
                 size_pt=16, bold=True)
        run_value = p.add_run(value)
        set_font(run_value, name_cn=font_cn, name_en="Times New Roman",
                 size_pt=16, bold=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Pt(0)

    doc.add_page_break()


# ─────────────────────────── 主解析逻辑 ────────────────────────────

def parse_table_line(line):
    """解析 | a | b | c | 格式的表格行，返回单元格列表。"""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_table_separator(line):
    """判断是否是表格分隔行（|:---|:---:|...）。"""
    stripped = line.strip().replace(' ', '')
    return bool(re.match(r'^[|\-:]+$', stripped)) and '|' in stripped and '-' in stripped


def convert_md_to_docx(md_path: str, docx_path: str):
    doc = Document()
    setup_page(doc)
    add_header(doc)
    add_page_number(doc)

    lines = Path(md_path).read_text(encoding='utf-8').splitlines()
    md_dir = Path(md_path).resolve().parent

    cover_info = extract_cover_info(lines)
    output_path = Path(docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 状态机
    in_code_block = False
    code_lang = ''
    code_lines = []
    pending_diagram_caption = None
    in_table = False
    table_header = []
    table_data_rows = []
    ordered_counters = {}   # indent_level -> count
    skip_cover_section = True
    skip_toc_body = False
    level1_count = 0
    current_section = ""

    # 生成封面
    add_title_page(
        doc,
        title=cover_info["题目"],
        title_en=cover_info["英文题目"],
        author=cover_info["学生姓名"],
        student_id=cover_info["学号"],
        advisor=cover_info["指导教师"],
        major=cover_info["专业名称"],
        college=cover_info["所在学院"],
        date=cover_info["完成日期"],
        print_date=cover_info["打印日期"],
    )

    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line
        stripped = line.strip()

        if skip_cover_section:
            if re.match(r'^#\s+原创性声明与使用授权', stripped):
                skip_cover_section = False
            else:
                i += 1
                continue

        if skip_toc_body and not re.match(r'^#\s+', stripped):
            i += 1
            continue
        if skip_toc_body and re.match(r'^#\s+', stripped):
            skip_toc_body = False

        # ── 代码块 ──────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                if code_lang.lower() in {'mermaid', 'plantuml'}:
                    pending_diagram_caption = "待渲染流程图"
                else:
                    add_code_block(doc, '\n'.join(code_lines))
                    doc.add_paragraph()  # 代码块后空行
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── 标题 ─────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            # 如果正在收集表格，先输出
            if in_table and table_header:
                add_table(doc, table_header, table_data_rows)
                in_table = False; table_header = []; table_data_rows = []

            level = len(heading_match.group(1))
            text = normalize_heading_text(heading_match.group(2))

            if level == 1:
                if level1_count > 0:
                    doc.add_page_break()
                level1_count += 1
                current_section = text

            # 目录标题直接转为提示页，正文目录项不再逐行写入
            if level == 1 and text == '目  录':
                add_heading(doc, text, 1)
                p = doc.add_paragraph()
                run = p.add_run('（此处请在 Word 中通过"引用 → 目录"自动生成目录）')
                set_font(run, name_cn="宋体", name_en="Times New Roman",
                         size_pt=12)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = Pt(20)
                skip_toc_body = True
                i += 1
                continue

            add_heading(doc, text, level)
            i += 1
            continue

        # ── 分隔线 ───────────────────────────────────────
        if re.match(r'^---+\s*$', line.strip()):
            # 如果正在收集表格，先输出
            if in_table and table_header:
                add_table(doc, table_header, table_data_rows)
                in_table = False; table_header = []; table_data_rows = []
            i += 1
            continue

        # ── 表格行 ───────────────────────────────────────
        if re.match(r'^\|', line.strip()):
            if is_table_separator(line):
                i += 1
                continue
            cells = parse_table_line(line)
            if not in_table:
                # 第一行是表头
                in_table = True
                table_header = cells
                table_data_rows = []
            else:
                table_data_rows.append(cells)
            i += 1
            continue
        else:
            # 非表格行：如果之前在收集表格，优先识别紧随其后的表题并置于表格上方
            if in_table and table_header:
                table_caption = extract_table_caption(stripped)
                if table_caption:
                    add_caption(doc, table_caption, before_pt=12, after_pt=6, caption_type="table")
                    add_table(doc, table_header, table_data_rows)
                    in_table = False; table_header = []; table_data_rows = []
                    i += 1
                    continue
                add_table(doc, table_header, table_data_rows)
                in_table = False; table_header = []; table_data_rows = []

        # ── 空行 ─────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Markdown 图片：![图4.1 单任务处理页面](materials/screenshots/a.png)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', stripped)
        if img_match:
            caption_text = img_match.group(1).strip() or "图片"
            image_ref = img_match.group(2).strip().strip('"')
            image_path = Path(image_ref)
            if not image_path.is_absolute():
                image_path = md_dir / image_path
            add_image_with_caption(doc, image_path, caption_text)
            i += 1
            continue

        # ── 占位图注释块 ─────────────────────────────────
        # 格式: **[占位图 X-X]** 或 [占位图 X-X]
        ph_match = re.match(r'.*\[占位图\s+([^\]]+)\]', stripped)
        if ph_match:
            caption_id = ph_match.group(1)
            # 向后找 > *图X-X 说明行*
            caption_text = f"占位图 {caption_id}"
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                desc_match = re.match(r'>\s*\*?(图\d+-\d+.*?)\*?', lines[j].strip())
                if desc_match:
                    caption_text = desc_match.group(1).strip().rstrip('）').strip()
                    i = j + 1
                else:
                    i += 1
            else:
                i += 1
            add_placeholder_box(doc, caption_text)
            add_caption(doc, f"（占位：{caption_text}）")
            continue

        figure_caption = extract_figure_caption(stripped)
        if figure_caption:
            pending_diagram_caption = None
            add_placeholder_box(doc, figure_caption)
            add_caption(doc, figure_caption)
            i += 1
            continue

        if pending_diagram_caption and re.match(r'^图\s*\d+(?:[.-]\d+)+', stripped):
            add_placeholder_box(doc, stripped)
            add_caption(doc, stripped)
            pending_diagram_caption = None
            i += 1
            continue

        # 跳过 > *图... 行（已被占位图处理消耗）
        if re.match(r'^>\s*\*?图\d+(?:[.-]\d+)+', stripped):
            i += 1
            continue

        # ── blockquote 说明行（非图注） ────────────────────
        if stripped.startswith('>'):
            text_content = stripped.lstrip('> ').strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text_content)
            set_font(run, name_cn="宋体", name_en="Times New Roman",
                     size_pt=12)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── 公式行（含 $$ 或 \tag 或 $$...$$） ──────────────
        if '$$' in stripped or r'\tag' in stripped or re.match(r'^\$\$', stripped):
            add_formula_line(doc, stripped)
            i += 1
            continue

        # ── 有序列表 ─────────────────────────────────────
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ol_match:
            indent = len(ol_match.group(1)) // 2
            number = int(ol_match.group(2))
            text_content = ol_match.group(3)
            add_list_item(doc, text_content, ordered=True,
                          number=number, level=indent)
            i += 1
            continue

        # ── 无序列表（- 或 * 开头，非表格分隔） ──────────────
        ul_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if ul_match:
            indent = len(ul_match.group(1)) // 2
            text_content = ul_match.group(2)
            # 跳过纯 --- 或 * 的分隔行（已在上面处理）
            add_list_item(doc, text_content, ordered=False, level=indent)
            i += 1
            continue

        # ── 表格标题行（**表X-X ...**） ───────────────────
        table_caption = extract_table_caption(stripped)
        if table_caption:
            add_caption(doc, table_caption, before_pt=12, after_pt=6, caption_type="table")
            i += 1
            continue

        # ── 诚信声明分页 ─────────────────────────────────
        if '诚信声明' in stripped and stripped.startswith('#'):
            doc.add_page_break()

        # ── 普通正文段落 ─────────────────────────────────
        # 判断是否是摘要/关键词行（无缩进）
        no_indent_prefixes = ('**关键词', '**Keywords', '关键词：', 'Key words:',
                              '声明人签名', '学生签名', '指导教师签名', '日期：',
                              'Co-Authored', '作者签名', '\\*热红外')
        is_no_indent = any(stripped.startswith(p) for p in no_indent_prefixes)
        first_indent = 0 if is_no_indent else 2

        add_special_paragraph(doc, stripped, section_name=current_section, first_indent=first_indent)
        i += 1

    # 收尾：若表格还未提交
    if in_table and table_header:
        add_table(doc, table_header, table_data_rows)

    doc.save(output_path)
    print(f"[OK] 已生成: {output_path}")


# ─────────────────────────── 入口 ────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        src = sys.argv[1]
        dst = sys.argv[2]
    else:
        # 默认路径
        base = Path(__file__).parent.parent
        src = str(base / 'docs' / '毕业设计论文_v1.md')
        dst = str(base / 'docs' / '毕业设计论文_v1.docx')

    print(f"源文件 : {src}")
    print(f"目标文件: {dst}")
    convert_md_to_docx(src, dst)
