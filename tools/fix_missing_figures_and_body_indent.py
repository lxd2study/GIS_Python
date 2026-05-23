from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "doc" / "基于Web的Landsat8遥感影像在线预处理系统-内容丰富版.docx"
ASSET_DIR = ROOT / "output" / "doc" / "thesis_generated_figures"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size=size)
    return ImageFont.load_default()


FONT_TITLE = font(38, True)
FONT_HEAD = font(25, True)
FONT_BODY = font(21)
FONT_SMALL = font(18)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt=FONT_BODY, fill="#273849") -> None:
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 28)
    line_h = fnt.size + 8
    y = y1 + (y2 - y1 - line_h * len(lines)) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1500, 900), "#f7f9fb")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((34, 34, 1466, 866), radius=24, fill="#ffffff", outline="#d5e1eb", width=2)
    draw.text((78, 62), title, font=FONT_TITLE, fill="#183247")
    return img, draw


def box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, fill: str, outline: str = "#6c8ebf") -> None:
    draw.rounded_rectangle(rect, radius=16, fill=fill, outline=outline, width=3)
    centered(draw, rect, text)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#627382") -> None:
    draw.line((start, end), fill=color, width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 18, ey - 10), (ex - sign * 18, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - sign * 18), (ex + 10, ey - sign * 18)]
    draw.polygon(pts, fill=color)


def make_flow(path: Path, title: str, nodes: list[str]) -> None:
    img, draw = canvas(title)
    y = 380
    gap = 28
    w = int((1340 - gap * (len(nodes) - 1)) / len(nodes))
    x = 80
    colors = ["#dae8fc", "#d5e8d4", "#fff2cc", "#f8cecc", "#e1d5e7", "#d5e8d4"]
    for i, text in enumerate(nodes):
        box(draw, (x, y, x + w, y + 120), text, colors[i % len(colors)])
        if i < len(nodes) - 1:
            arrow(draw, (x + w, y + 60), (x + w + gap, y + 60))
        x += w + gap
    img.save(path)


def make_grid(path: Path, title: str, center_text: str, items: list[str]) -> None:
    img, draw = canvas(title)
    box(draw, (575, 375, 925, 505), center_text, "#dae8fc")
    positions = [
        (120, 210, 410, 315),
        (605, 190, 895, 295),
        (1090, 210, 1380, 315),
        (120, 620, 410, 725),
        (605, 640, 895, 745),
        (1090, 620, 1380, 725),
    ]
    for idx, text in enumerate(items):
        rect = positions[idx]
        box(draw, rect, text, ["#d5e8d4", "#fff2cc", "#f8cecc", "#e1d5e7", "#dae8fc", "#fff2cc"][idx])
        cx = (rect[0] + rect[2]) // 2
        cy = (rect[1] + rect[3]) // 2
        arrow(draw, (cx, cy), (750, 440))
    img.save(path)


def make_use_case(path: Path) -> None:
    img, draw = canvas("系统功能用例图")
    draw.ellipse((110, 350, 220, 460), outline="#2f4858", width=4)
    draw.line((165, 460, 165, 610), fill="#2f4858", width=4)
    draw.line((85, 520, 245, 520), fill="#2f4858", width=4)
    draw.line((165, 610, 95, 725), fill="#2f4858", width=4)
    draw.line((165, 610, 235, 725), fill="#2f4858", width=4)
    centered(draw, (70, 735, 260, 790), "系统用户", FONT_HEAD)
    use_cases = [
        "单景预处理", "批量流程编排", "影像检索下载",
        "AOI 范围配置", "结果预览下载", "指数信息查询",
    ]
    positions = [(480, 170), (850, 170), (480, 350), (850, 350), (480, 530), (850, 530)]
    for text, (x, y) in zip(use_cases, positions):
        draw.ellipse((x, y, x + 270, y + 92), fill="#eef6ff", outline="#6c8ebf", width=3)
        centered(draw, (x, y, x + 270, y + 92), text)
        arrow(draw, (245, 520), (x, y + 46))
    img.save(path)


def generate_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "图1.2": ASSET_DIR / "fig-1-02-research-framework.png",
        "图2.1": ASSET_DIR / "fig-2-01-preprocess-flow.png",
        "图2.3": ASSET_DIR / "fig-2-03-tech-relationship.png",
        "图2.4": ASSET_DIR / "fig-2-04-use-case.png",
        "图2.5": ASSET_DIR / "fig-2-05-user-flow.png",
        "图3.2": ASSET_DIR / "fig-3-02-module-structure.png",
        "图3.3": ASSET_DIR / "fig-3-03-single-flow.png",
        "图3.4": ASSET_DIR / "fig-3-04-batch-flow.png",
        "图3.5": ASSET_DIR / "fig-3-05-download-flow.png",
        "图3.6": ASSET_DIR / "fig-3-06-result-flow.png",
        "图4.10": ASSET_DIR / "fig-4-10-l1-l2-logic.png",
        "图4.11": ASSET_DIR / "fig-4-11-mosaic-flow.png",
    }
    make_flow(assets["图1.2"], "论文研究内容框架", ["背景与现状", "关键技术", "总体设计", "详细实现", "测试评价", "总结展望"])
    make_flow(assets["图2.1"], "遥感影像预处理基本流程", ["数据输入", "辐射处理", "大气校正", "质量掩膜", "裁剪合成", "结果输出"])
    make_grid(assets["图2.3"], "系统关键开发技术关系图", "Web 遥感预处理系统", ["FastAPI 接口服务", "GDAL / NumPy 栅格处理", "Py6S / DOS 校正", "OpenLayers 空间交互", "Vue Flow 批量编排", "STAC 检索下载"])
    make_use_case(assets["图2.4"])
    make_flow(assets["图2.5"], "用户业务流程图", ["准备数据", "配置参数", "提交任务", "查看进度", "预览结果", "下载归档"])
    make_grid(assets["图3.2"], "系统功能模块结构图", "功能模块", ["单景预处理", "批量处理", "多景镶嵌", "影像检索下载", "结果资产中心", "指数信息辅助"])
    make_flow(assets["图3.3"], "单景预处理业务流程图", ["上传/选择波段", "解析 MTL 与 QA", "选择 L1/L2 链路", "裁剪与掩膜", "合成/指数", "写入清单"])
    make_flow(assets["图3.4"], "批量任务执行流程图", ["扫描场景", "构建节点图", "拓扑排序", "生成任务", "队列执行", "汇总结果"])
    make_flow(assets["图3.5"], "影像在线检索下载流程图", ["设置 AOI", "选择集合", "STAC 检索", "资产选择", "服务端下载", "本地归档"])
    make_flow(assets["图3.6"], "结果资产管理流程图", ["任务完成", "生成 manifest", "扫描历史目录", "分类产物", "预览/下载", "结果复用"])
    make_flow(assets["图4.10"], "L1/L2 双链路处理逻辑图", ["识别产品级别", "L1 定标校正", "L2 缩放换算", "质量掩膜", "合成指数", "结果归档"])
    make_flow(assets["图4.11"], "多景镶嵌处理流程图", ["逐景预处理", "收集中间波段", "同名波段镶嵌", "统一裁剪", "显示匀色", "输出合成图"])
    return assets


def insert_picture_before(paragraph: Paragraph, image_path: Path, width_cm: float = 14.2) -> None:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    pic_p = Paragraph(new_p, paragraph._parent)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.first_line_indent = Pt(0)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def previous_has_picture(doc: Document, index: int) -> bool:
    return index > 0 and bool(doc.paragraphs[index - 1]._p.xpath(".//w:drawing"))


def insert_missing_figures(doc: Document, assets: dict[str, Path]) -> list[str]:
    inserted: list[str] = []
    pattern = re.compile(r"^(图\d+\.\d+)\u3000")
    for index, paragraph in list(enumerate(doc.paragraphs)):
        text = paragraph.text.strip()
        match = pattern.match(text)
        if not match:
            continue
        key = match.group(1)
        if key in assets and not previous_has_picture(doc, index):
            insert_picture_before(paragraph, assets[key])
            inserted.append(key)
    return inserted


def normalize_body_indent(doc: Document) -> int:
    changed = 0
    in_main = False
    code_mode = False
    skip_prefixes = ("图", "表", "代码清单")
    front_titles = {
        "本科毕业设计", "毕业设计原创性声明", "毕业设计版权使用授权书",
        "摘　　要", "ABSTRACT", "目　　录", "参考文献", "致　　谢",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name
        if text.startswith("第1章"):
            in_main = True
        if text == "参考文献":
            in_main = False
        if not text:
            continue
        if text.startswith("代码清单"):
            code_mode = True
            continue
        if style.startswith("Heading") or style.startswith("toc") or text.startswith(skip_prefixes):
            code_mode = False if not text.startswith("代码清单") else code_mode
            continue
        if not in_main or code_mode or re.match(r"^\[?\d+\]?", text) and "http" in text:
            continue
        if text in front_titles:
            continue
        if paragraph.paragraph_format.first_line_indent != Pt(24):
            paragraph.paragraph_format.first_line_indent = Pt(24)
            changed += 1
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    return changed


def main() -> None:
    assets = generate_assets()
    doc = Document(DOCX)
    inserted = insert_missing_figures(doc, assets)
    changed = normalize_body_indent(doc)
    doc.save(DOCX)
    print(f"inserted_figures={inserted}")
    print(f"body_indent_changed={changed}")
    print(DOCX)


if __name__ == "__main__":
    main()
